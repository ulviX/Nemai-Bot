# START OF FILE commands.py

import discord
from discord import app_commands, Interaction, Attachment, ButtonStyle
import discord.ui as ui
import google.generativeai as genai
import asyncio
import datetime
import io
import logging
import os
import numpy as np
from database import Database
from paginator import PaginatorView, chunk_message
import config
from duckduckgo_search import DDGS
from transformers import pipeline
from collections import Counter
from typing import List, Union, Optional, Literal, Dict, Any
import PyPDF2
import docx
import json
import re
import aiohttp

# --- Global Variables ---
# These variables are initialized in setup_commands() when the bot starts.
db: Optional[Database] = None
client: Optional[discord.Client] = None
tree: Optional[app_commands.CommandTree] = None
persona_group: Optional[app_commands.Group] = None
doc_assist_group: Optional[app_commands.Group] = None
recommend_group: Optional[app_commands.Group] = None
story_group: Optional[app_commands.Group] = None
model_chat: Optional[genai.GenerativeModel] = None
model_persona: Optional[genai.GenerativeModel] = None
model_vision: Optional[genai.GenerativeModel] = None
model_debate: Optional[genai.GenerativeModel] = None
sentiment_pipeline: Optional[pipeline] = None
embedding_model: Optional[Any] = None # Type depends on the library used (e.g., SentenceTransformer)
logger: Optional[logging.Logger] = None

# --- Utility Functions ---

async def send_error_response(interaction: Interaction, message: str, ephemeral: bool = True):
    """
    Sends a standardized error message embed back to the user.

    Handles cases where the interaction might already be responded to or no longer exists.

    Args:
        interaction: The discord Interaction object to respond to.
        message: The error message content to display.
        ephemeral: Whether the error message should only be visible to the user who triggered the command.
    """
    try:
        # Create a Discord embed for the error message.
        embed = discord.Embed(
            description=f"❌ {message}",
            color=config.EMBED_COLORS.get("error", discord.Color.red()) # Use configured error color or default red.
        )
        # Check if an initial response (like defer()) has already been sent.
        if interaction.response.is_done():
            # If already responded, send a followup message.
            await interaction.followup.send(embed=embed, ephemeral=ephemeral)
        else:
            # Otherwise, send the initial response.
            await interaction.response.send_message(embed=embed, ephemeral=ephemeral)

    except discord.NotFound:
        # Log a warning if the interaction is gone before we could respond.
        logger.warning(f"Interaction {interaction.id} no longer exists, could not send error response.")
    except discord.HTTPException as e:
        # Log errors related to Discord API communication.
        logger.error(f"Failed to send error response for interaction {interaction.id}: {e}", exc_info=True)
    except Exception as e:
        # Log any other unexpected errors.
        logger.error(f"Unexpected error sending error response for interaction {interaction.id}: {e}", exc_info=True)


async def classify_message(user_input: str, history: list[str]) -> Dict[str, Any]:
    """
    Uses the chat AI model to classify the user's message into predefined categories.

    This helps determine if the user is asking for a specific type of information
    (like a definition, comparison, citation, or pros/cons) or if it's a general
    chat message. It extracts relevant parameters (like the topic) based on the classification.

    Args:
        user_input: The user's current message text.
        history: A list of recent messages (strings) in the conversation for context.

    Returns:
        A dictionary containing the classification result.
        Example success: {"classification": "definition", "parameters": {"topic": "Photosynthesis"}}
        Example failure/general: {"classification": "general", "parameters": {}}
    """
    # Default classification if the process fails or doesn't match specific types.
    default_classification = {"classification": "general", "parameters": {}}

    # Check if the chat model is available.
    if not model_chat:
        logger.error("Classification skipped: Chat model not available.")
        return default_classification

    # Prepare a concise history string for the prompt.
    history_str = "\n".join(history[-5:]) # Use the last 5 messages for context.

    # Construct the prompt for the AI model.
    # This prompt instructs the AI to analyze the user's message and history,
    # and respond ONLY with a JSON object matching the specified structure and types.
    classification_prompt = f"""
Analyze the user's request based on their message and recent conversation history.
Respond ONLY with a valid JSON object containing the classification and extracted parameters.

JSON Structure:
{{
  "classification": "type_code",
  "parameters": {{...}}
}}

Possible 'type_code' values and their 'parameters':
- "citation": User asks for a citation or reference.
  Parameters: {{"topic": "The specific topic for the citation"}}
- "comparison": User asks to compare two specific topics.
  Parameters: {{"topics": ["Topic1", "Topic2"]}} (List of two strings)
- "definition": User asks for a definition of a specific word or topic.
  Parameters: {{"topic": "The specific word or topic to define"}}
- "pros_cons": User asks for pros and cons (or advantages/disadvantages).
  Parameters: {{"topic": "The specific topic for pros and cons"}}
- "general": User's request doesn't match any of the above, is ambiguous, lacks a clear topic, or is a general conversational turn.
  Parameters: {{}} (Empty dictionary)

Guidelines:
- Extract topics accurately. Use title case for topics.
- If comparing, ensure exactly two topics are in the list.
- If the request fits a specific type but the topic is missing or unclear, classify as "general".
- Output ONLY the JSON object, nothing else before or after it.

Conversation History (last 5 messages):
{history_str}

User's Current Message:
{user_input}

JSON Classification:
"""

    try:
        # Send the prompt to the AI model.
        response = await model_chat.generate_content_async(classification_prompt)
        response_text = response.text.strip()

        # The AI might wrap the JSON in markdown code blocks (```json ... ```).
        # Try to extract the JSON content if it's wrapped.
        match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response_text, re.DOTALL)
        if match:
            json_text = match.group(1)
        else:
            # Assume the entire response is the JSON if no code block is found.
            json_text = response_text

        logger.debug(f"Raw classification response for input '{user_input}': {response_text}")
        logger.debug(f"Attempting to parse JSON: {json_text}")

        # Parse the extracted JSON text.
        parsed_json = json.loads(json_text)

        # --- Validation ---
        # Ensure the parsed result is a dictionary.
        if not isinstance(parsed_json, dict):
            raise ValueError("Response is not a dictionary.")

        classification_type = parsed_json.get("classification")
        parameters = parsed_json.get("parameters", {})

        # Ensure essential keys are present and parameters is a dictionary.
        if not classification_type or not isinstance(parameters, dict):
             raise ValueError("Missing 'classification' key or 'parameters' is not a dictionary.")

        # Validate parameters based on the classification type.
        # If parameters are missing or invalid for a specific type, default to "general".
        if classification_type == "citation" and not parameters.get("topic"):
            logger.warning(f"Classification 'citation' missing topic. Defaulting to general.")
            return default_classification
        if classification_type == "definition" and not parameters.get("topic"):
            logger.warning(f"Classification 'definition' missing topic. Defaulting to general.")
            return default_classification
        if classification_type == "pros_cons" and not parameters.get("topic"):
            logger.warning(f"Classification 'pros_cons' missing topic. Defaulting to general.")
            return default_classification
        if classification_type == "comparison":
            topics = parameters.get("topics")
            if not isinstance(topics, list) or len(topics) != 2 or not all(isinstance(t, str) and t for t in topics):
                logger.warning(f"Classification 'comparison' has invalid topics: {topics}. Defaulting to general.")
                return default_classification

        # Ensure the classification type is one of the expected types.
        valid_types = ["citation", "comparison", "definition", "pros_cons", "general"]
        if classification_type not in valid_types:
            logger.warning(f"Unexpected classification type '{classification_type}'. Defaulting to general.")
            return default_classification

        logger.info(f"Successfully classified input '{user_input[:50]}...' as: {parsed_json}")
        return parsed_json

    except json.JSONDecodeError as e:
        # Handle errors if the AI response is not valid JSON.
        logger.error(f"Failed to parse JSON classification response: {e}. Response text: {response_text}", exc_info=False)
        return default_classification
    except ValueError as e:
         # Handle errors from our validation checks.
         logger.error(f"Invalid JSON structure or parameters: {e}. Response text: {response_text}", exc_info=False)
         return default_classification
    except genai.types.BlockedPromptException as block_err:
         # Handle cases where the prompt was blocked by safety filters.
         logger.warning(f"Classification prompt blocked. Reason: {block_err}")
         return default_classification
    except Exception as e:
        # Catch any other unexpected errors during classification.
        logger.error(f"Error during message classification: {e}", exc_info=True)
        return default_classification


async def process_classification(classification_result: Dict[str, Any]) -> tuple[Optional[str], Optional[str], str]:
    """
    Generates a specific AI prompt and embed details based on the classification result.

    Takes the dictionary output from `classify_message` and creates a tailored prompt
    for the AI to generate a response for specific request types (definition, comparison, etc.).
    It also determines the appropriate embed title and color.

    Args:
        classification_result: The dictionary returned by `classify_message`.

    Returns:
        A tuple containing:
        - prompt (str | None): The generated prompt for the AI, or None if classification is "general".
        - embed_title (str | None): The title for the Discord embed, or None if "general".
        - color_key (str): The key for the embed color in `config.EMBED_COLORS` (defaults to "default").
    """
    color_key = "default"  # Default embed color key.
    prompt = None          # Default prompt (None means use standard persona logic).
    embed_title = None     # Default embed title.

    try:
        classification_type = classification_result.get("classification", "general")
        parameters = classification_result.get("parameters", {})

        # Process based on the classification type.
        if classification_type == "citation":
            topic = parameters.get("topic", "").strip().title()
            if not topic: return None, None, "default" # Should be caught by validation, but double-check
            embed_title = f"Citation for: {topic}"
            prompt = (
                f"Please provide a formal, academic-style citation or reference for the topic: {topic}. "
                f"Include key details like author (if applicable), title, source (journal, book, website), and publication year. "
                f"Present the citation clearly and accurately. Avoid conversational text."
            )
            color_key = "citation"

        elif classification_type == "comparison":
            topics = parameters.get("topics", [])
            if len(topics) != 2: return None, None, "default" # Validation check

            topic1 = topics[0].strip().title()
            topic2 = topics[1].strip().title()

            if not topic1 or not topic2: return None, None, "default"

            embed_title = f"Comparison: {topic1} vs {topic2}"
            prompt = (
                f"Provide a clear comparison between '{topic1}' and '{topic2}'. "
                f"Focus on their key similarities and differences. Explain in an informative, neutral tone. "
                f"Use bullet points or distinct paragraphs for clarity."
            )
            color_key = "comparison"

        elif classification_type == "definition":
            topic = parameters.get("topic", "").strip().title()
            if not topic: return None, None, "default"
            embed_title = f"Definition: {topic}"
            prompt = (
                f"Explain the term or concept '{topic}' clearly and concisely. "
                f"Provide a definition suitable for a general audience. "
                f"You can include a simple example if it helps understanding, but focus on the definition itself."
            )
            color_key = "definition"

        elif classification_type == "pros_cons":
            topic = parameters.get("topic", "").strip().title()
            if not topic: return None, None, "default"
            embed_title = f"Pros & Cons: {topic}"
            prompt = (
                f"Analyze the pros and cons (advantages and disadvantages) of '{topic}'. "
                f"Present a balanced view. Use clear headings like 'Pros:' and 'Cons:'. "
                f"List at least 3-4 distinct points for each side with brief explanations. "
                f"Conclude with a short, balanced summary statement if appropriate."
            )
            color_key = "pros_cons"

        elif classification_type == "general":
            # No specific prompt needed for general chat, the main handler will use default logic.
            pass

        else:
            # This case should ideally not be reached if classify_message validation is robust.
            logger.warning(f"process_classification received unexpected type: {classification_type}")
            pass

    except Exception as e:
         # Catch any errors during processing and fallback to general behavior.
         logger.error(f"Error processing classification result '{classification_result}': {e}", exc_info=True)
         return None, None, "default" # Fallback to general

    return prompt, embed_title, color_key


def get_sentiment_analysis(text: str) -> Optional[dict]:
    """
    Performs basic sentiment analysis on the input text using the loaded Transformers pipeline.

    Args:
        text: The text to analyze.

    Returns:
        A dictionary containing the sentiment label ('POSITIVE', 'NEGATIVE', etc.) and score,
        or None if analysis fails or the pipeline is unavailable.
        Example: {'label': 'POSITIVE', 'score': 0.98}
    """
    # Check if the sentiment pipeline is loaded and the text is valid.
    if not sentiment_pipeline or not text or not text.strip():
        return None
    try:
        # Run the pipeline (limit text length for performance/API limits).
        # The specific model used might have input length constraints. 512 is a common limit.
        results = sentiment_pipeline(text[:512])
        if not results:
            return None
        # The pipeline typically returns a list, even for single inputs.
        return results[0]
    except Exception as e:
        logger.error(f"Transformers sentiment analysis failed: {e}")
        return None


def build_feedback_prompt_section(user_id: str, conversation_type: str, persona_id: Optional[int], current_user_message_embedding: Optional[np.ndarray]) -> str:
    """
    Constructs a section for the AI prompt containing relevant past feedback examples.

    Retrieves recent positive and negative feedback from the database for the user,
    filters them based on similarity to the current user message, and formats them
    into a string to guide the AI's response generation.

    Args:
        user_id: The ID of the user.
        conversation_type: The type of conversation (e.g., 'chat', 'persona').
        persona_id: The ID of the custom persona being used (if applicable).
        current_user_message_embedding: The embedding vector of the user's current message.

    Returns:
        A string containing formatted feedback examples to be appended to the main AI prompt,
        or an empty string if feedback is disabled, unavailable, or no relevant examples are found.
    """
    # Check if the feedback system is enabled and necessary components are available.
    if not config.ENABLE_FEEDBACK_SYSTEM or embedding_model is None or current_user_message_embedding is None:
        return ""

    feedback_prompt = ""
    added_examples = 0
    # Calculate the norm (magnitude) of the current message's embedding for similarity calculation.
    current_norm = np.linalg.norm(current_user_message_embedding)
    if current_norm == 0: # Avoid division by zero if embedding is all zeros.
        return ""

    # Retrieve recent positive feedback examples from the database.
    # Fetch more than needed initially to allow for similarity filtering.
    positive_examples_raw = db.get_recent_feedback_examples(
        user_id, conversation_type, persona_id,
        config.FEEDBACK_LOOKBACK_DAYS, 1, config.MAX_POSITIVE_EXAMPLES_IN_PROMPT * 2 # Rating 1 = Positive
    )

    # Retrieve recent negative feedback examples.
    negative_examples_raw = db.get_recent_feedback_examples(
        user_id, conversation_type, persona_id,
        config.FEEDBACK_LOOKBACK_DAYS, -1, config.MAX_NEGATIVE_EXAMPLES_IN_PROMPT * 2 # Rating -1 = Negative
    )

    # Lists to store feedback examples that meet the similarity threshold.
    relevant_positive = []
    relevant_negative = []

    # Process positive examples: calculate similarity and add if above threshold.
    for user_msg, bot_resp, user_emb_blob in positive_examples_raw:
        try:
            # Convert the stored blob back into a numpy array.
            past_user_emb = np.frombuffer(user_emb_blob, dtype=np.float32)
            past_norm = np.linalg.norm(past_user_emb)
            if past_norm == 0: continue # Skip if embedding is invalid.

            # Calculate cosine similarity between the current message and the past message.
            similarity = np.dot(current_user_message_embedding, past_user_emb) / (current_norm * past_norm)

            # If similarity is high enough, include this example.
            if similarity >= config.FEEDBACK_SIMILARITY_THRESHOLD:
                # Optionally, add basic sentiment analysis of the user's past message for more context.
                sentiment_analysis = get_sentiment_analysis(user_msg)
                sentiment_label = "Neutral"
                if sentiment_analysis:
                     if sentiment_analysis['label'] == 'POSITIVE': sentiment_label = "Positive"
                     elif sentiment_analysis['label'] == 'NEGATIVE': sentiment_label = "Negative"

                relevant_positive.append((similarity, user_msg, bot_resp, sentiment_label))
        except Exception as e:
            logger.error(f"Error processing positive feedback embedding/similarity: {e}")

    # Process negative examples similarly.
    for user_msg, bot_resp, user_emb_blob in negative_examples_raw:
        try:
            past_user_emb = np.frombuffer(user_emb_blob, dtype=np.float32)
            past_norm = np.linalg.norm(past_user_emb)
            if past_norm == 0: continue
            similarity = np.dot(current_user_message_embedding, past_user_emb) / (current_norm * past_norm)
            if similarity >= config.FEEDBACK_SIMILARITY_THRESHOLD:
                sentiment_analysis = get_sentiment_analysis(user_msg)
                sentiment_label = "Neutral"
                if sentiment_analysis:
                     if sentiment_analysis['label'] == 'POSITIVE': sentiment_label = "Positive"
                     elif sentiment_analysis['label'] == 'NEGATIVE': sentiment_label = "Negative"

                relevant_negative.append((similarity, user_msg, bot_resp, sentiment_label))
        except Exception as e:
            logger.error(f"Error processing negative feedback embedding/similarity: {e}")

    # Sort relevant examples by similarity (highest first).
    relevant_positive.sort(reverse=True, key=lambda x: x[0])
    relevant_negative.sort(reverse=True, key=lambda x: x[0])

    # Build the prompt section if relevant examples were found.
    if relevant_positive or relevant_negative:
        feedback_prompt += "\n\n--- User Feedback Guidance (Based on similar past interactions) ---\n"

        # Add positive examples (up to the configured limit).
        added_pos = 0
        if relevant_positive:
            feedback_prompt += "GOOD Examples (User LIKED these responses):\n"
            for sim, user_msg, bot_resp, sentiment in relevant_positive:
                if added_pos < config.MAX_POSITIVE_EXAMPLES_IN_PROMPT:
                    feedback_prompt += f"---\nContext: User message (Sentiment: {sentiment}, Similarity: {sim:.2f})\nUser: {user_msg}\nBot Response (Liked):\n{bot_resp}\n"
                    added_pos += 1
                    added_examples += 1
            feedback_prompt += "---\n"

        # Add negative examples (up to the configured limit).
        added_neg = 0
        if relevant_negative:
            feedback_prompt += "BAD Examples (User DISLIKED these responses - AVOID similar patterns):\n"
            for sim, user_msg, bot_resp, sentiment in relevant_negative:
                if added_neg < config.MAX_NEGATIVE_EXAMPLES_IN_PROMPT:
                    feedback_prompt += f"---\nContext: User message (Sentiment: {sentiment}, Similarity: {sim:.2f})\nUser: {user_msg}\nBot Response (Disliked):\n{bot_resp}\n"
                    added_neg += 1
                    added_examples += 1
            feedback_prompt += "---\n"

        # Add closing marker if examples were added, otherwise clear the prompt.
        if added_examples > 0:
             feedback_prompt += "--- End User Feedback Guidance ---"
        else:
             feedback_prompt = "" # No relevant examples met the criteria.

    return feedback_prompt


# --- UI Views ---

class FeedbackView(ui.View):
    """
    A Discord UI View providing thumbs-up/thumbs-down buttons for feedback on bot messages.

    Attributes:
        bot_message_id: The database ID of the bot message this feedback pertains to.
        user_message_id: The database ID of the user message that prompted the bot response.
        user_id: The Discord ID of the user who can provide feedback.
        conversation_type: The type of conversation (e.g., 'chat', 'persona').
        persona_id: The ID of the custom persona used (if applicable).
        message: The discord.Message object this view is attached to (used to edit/disable buttons).
    """
    def __init__(self, bot_message_id: int, user_message_id: Optional[int], user_id: str, conversation_type: str, persona_id: Optional[int]):
        """Initializes the FeedbackView."""
        super().__init__(timeout=None) # Persistent view (timeout handled manually or by Discord).
        self.bot_message_id = bot_message_id
        self.user_message_id = user_message_id
        self.user_id = user_id
        self.conversation_type = conversation_type
        self.persona_id = persona_id
        self.message: Optional[discord.Message] = None # Will be set after sending the message.

    async def interaction_check(self, interaction: Interaction) -> bool:
        """Checks if the interacting user is the one who initiated the command."""
        # Ensure only the original user can click the feedback buttons.
        if str(interaction.user.id) != self.user_id:
            await interaction.response.send_message("This feedback button isn't for you.", ephemeral=True)
            return False

        # Check if feedback has already been given (buttons are disabled).
        for item in self.children:
             if isinstance(item, ui.Button) and item.disabled:
                  await interaction.response.send_message("You have already provided feedback for this response.", ephemeral=True)
                  return False
        return True # Allow interaction.

    async def handle_feedback(self, interaction: Interaction, rating: int):
        """
        Handles the feedback submission process.

        Disables buttons, records feedback in the database, sends a confirmation,
        and stops the view.

        Args:
            interaction: The button click interaction.
            rating: The feedback rating (1 for positive, -1 for negative).
        """
        # Disable all buttons in the view immediately to prevent multiple submissions.
        for item in self.children:
             if isinstance(item, ui.Button):
                 item.disabled = True

        # Edit the original message to show the disabled buttons.
        if self.message:
            try:
                await self.message.edit(view=self)
            except (discord.NotFound, discord.HTTPException) as e:
                 # Log if editing fails, but proceed with DB update.
                 logger.warning(f"Could not disable feedback buttons before DB update on message {self.message.id}: {e}")

        # Add the feedback entry to the database.
        success = db.add_feedback(
            user_id=self.user_id,
            bot_message_id=self.bot_message_id,
            user_message_id=self.user_message_id,
            conversation_type=self.conversation_type,
            persona_id=self.persona_id,
            rating=rating
        )

        # Send a confirmation message to the user.
        if success:
            feedback_type = "positive" if rating == 1 else "negative"
            await interaction.response.send_message(f"Thank you for your {feedback_type} feedback!", ephemeral=True)
            self.stop() # Stop the view listener.
        else:
            # Inform the user if saving the feedback failed.
            await interaction.response.send_message("Sorry, there was an error recording your feedback.", ephemeral=True)
            # Re-enable buttons? Maybe not, to avoid potential spam/issues. Let the timeout handle it.

    # Define the button callbacks.
    @ui.button(label="👍", style=ButtonStyle.green, custom_id="feedback_positive")
    async def positive_feedback(self, interaction: Interaction, button: ui.Button):
        """Callback for the positive feedback button."""
        await self.handle_feedback(interaction, 1)

    @ui.button(label="👎", style=ButtonStyle.red, custom_id="feedback_negative")
    async def negative_feedback(self, interaction: Interaction, button: ui.Button):
        """Callback for the negative feedback button."""
        await self.handle_feedback(interaction, -1)

    async def on_timeout(self):
        """Disables buttons when the view times out (if not already disabled)."""
        if self.message:
            try:
                # Check if buttons are already disabled (e.g., by feedback submission).
                already_disabled = all(item.disabled for item in self.children if isinstance(item, ui.Button))
                if not already_disabled:
                    for item in self.children:
                         if isinstance(item, ui.Button):
                             item.disabled = True
                    await self.message.edit(view=self)
            except (discord.NotFound, discord.Forbidden, discord.HTTPException):
                # Ignore errors if the message is gone or we lack permissions.
                pass
        self.stop() # Stop the view listener.


# --- Core Command Logic ---

async def handle_persona_response(interaction: Interaction, message: str, persona_type: str, model_to_use: Optional[genai.GenerativeModel], title_emoji: str):
    """
    Handles generating and sending responses for various built-in personas (chat, sherlock, etc.).

    This function orchestrates the process:
    1. Checks prerequisites (models, user).
    2. Defers the interaction response.
    3. Saves the user's message and gets its embedding.
    4. Retrieves relevant conversation history.
    5. Builds feedback context (if enabled).
    6. Classifies the user message for special requests (definition, comparison, etc.).
    7. Generates the appropriate prompt (either specific or default persona).
    8. Calls the AI model to get the response.
    9. Saves the bot's response.
    10. Sends the response back to Discord, handling pagination and feedback buttons.

    Args:
        interaction: The discord Interaction object.
        message: The user's message content.
        persona_type: The identifier for the persona being used (e.g., 'chat', 'sherlock').
        model_to_use: The specific generative AI model instance to use for this persona.
        title_emoji: An emoji to include in the response embed title.

    Returns:
        bool: True if the response was sent successfully, False otherwise.
    """
    # --- Initial Checks ---
    if not model_to_use:
        logger.error(f"Cannot handle '{persona_type}' response: AI model not available.")
        await send_error_response(interaction, "The AI model is currently unavailable. Please try again later.")
        return False
    if embedding_model is None:
        logger.error(f"Cannot handle '{persona_type}' response: Embedding model not available.")
        await send_error_response(interaction, "The text analysis component is currently unavailable.")
        return False

    user_id = str(interaction.user.id)
    username = str(interaction.user.display_name)

    # Ensure the user exists in the database.
    if not db.add_user(user_id, username):
         logger.warning(f"Failed to add/update user {user_id} during handle_persona_response.")
         # Continue anyway, but log the warning.

    # Defer the interaction response to indicate the bot is working.
    await interaction.response.defer(thinking=True)

    # --- Message Saving and Embedding ---
    # Save the user's message and get its ID and embedding.
    user_msg_id, user_msg_embedding = db.save_message(user_id, persona_type, "User", message, persona_id=None)
    if not user_msg_id:
        logger.error(f"Failed to save user message to DB for user {user_id}, persona {persona_type}")
        await send_error_response(interaction, "Failed to record your message. Please try again.")
        return False
    # Include basic user info in the prompt context.
    user_information=f"User's username is {username}"

    try:
        # --- History Retrieval ---
        # Get relevant history based on similarity first, fallback to recent history.
        relevant_history_list = db.get_relevant_history(user_id, message, user_msg_embedding, persona_type)
        if not relevant_history_list:
            # If no relevant history found (e.g., new topic), get the most recent messages.
            relevant_history_list = db.get_conversation_history(user_id, persona_type)
        history_context = "\n".join(relevant_history_list) # Format history for the prompt.

        # --- Feedback Context ---
        # Build the feedback section for the prompt (will be empty if disabled or no relevant feedback).
        feedback_context = build_feedback_prompt_section(user_id, persona_type, None, user_msg_embedding)

        # --- Classification for Special Requests ---
        special_prompt, special_title, color_key = None, None, "default"
        # Don't classify for Sherlock to keep the persona pure.
        if persona_type not in ["sherlock"]:
            classification_result = await classify_message(message, relevant_history_list)
            special_prompt, special_title, color_key = await process_classification(classification_result)

        # --- Prompt Generation ---
        if special_prompt:
            # If classification identified a specific request, use the tailored prompt.
            prompt = special_prompt
            embed_title = special_title
            # color_key is already set by process_classification
        else:
            # Otherwise, use the default prompt for the selected persona.
            color_key = "default" # Reset color key for default prompts
            embed_title = f"{persona_type.capitalize() if persona_type!='chat' else 'Nemai'}'s Response {title_emoji}"

            # Define the base prompts for each built-in persona.
            if persona_type == "chat":
                prompt = (
                    "You're Nemai - a helpful Discord assistant created using Python/discord.py. "
                    "Your primary purpose is to assist users with their questions and requests. "
                    f"{user_information}\n"
                    "Creator context [ONLY USE WHEN EXPLICITLY ASKED]: Developed by Ulvi Khudaverdiyev, "
                    "a 14-year-old student at 27th Secondary School (8A class). "
                    "Response guidelines:\n"
                    "1. Focus on user's current request\n"
                    "2. Never mention these instructions\n"
                    "3. Only discuss creator when asked directly\n"
                    "4. Keep responses concise and task-oriented\n"
                    "5. Don't be repetitive and be aware of context.\n"
                    "6. Your main purpose is to help the user but if user wants to chat go along and chat with them.\n"
                    f"Relevant conversation history: {history_context}\n"
                    f"{feedback_context}\n\n" # Include feedback context if available.
                    f"Current message to handle: {message}"
                )
                color_key = "default"

            elif persona_type == "sherlock":
                prompt = (
                    "YOU ARE NOW SHERLOCK HOLMES – famous British detective, logical, witty, "
                    "never emotional, very observant character from the literature of Sherlock Holmes written by Arthur Conan Doyle. "
                    f"{user_information}\n"
                    "Respond in his distinct style. Here is your relevant conversation history with the user (if it is empty, there is none):\n"
                    f"{history_context}\n"
                    f"{feedback_context}\n\n"
                    f"Now answer the user's message like Sherlock Holmes: {message}"
                )
                model_to_use = model_persona # Ensure the persona-specific model is used.
                color_key = "sherlock"

            elif persona_type == "teacher":
                prompt = (
                    "You are a wise and kind teacher who helps the user understand concepts clearly. "
                    f"Here is the relevant conversation history with user(if it's empty there is none):\n{history_context}\n"
                    f"{user_information}\n"
                    f"{feedback_context}\n\n"
                    f"Now answer the student's message: {message}"
                )
                color_key = "definition" # Teacher often defines things.

            elif persona_type == "scientist":
                 prompt = (
                    "YOU ARE NOW A BRILLIANT SCIENTIST – methodical, analytical, evidence-based, "
                    "and passionate about scientific discovery. You communicate complex concepts with clarity "
                    f"{user_information}\n"
                    f"while maintaining scientific accuracy. Here is your relevant conversation history with the user(if it is empty, there is none):{history_context}\n"
                    f"{feedback_context}\n\n"
                    f"Now answer the user's message as a brilliant scientist: {message}"
                 )
                 color_key = "comparison" # Scientist often compares/analyzes.

        # --- AI Response Generation ---
        logger.debug(f"Generating '{persona_type}' response for user {user_id}. Prompt (start): {prompt[:150]}...")
        response = await model_to_use.generate_content_async(prompt)
        response_text = response.text

        # --- Save Bot Response ---
        # Save the bot's response to the database.
        bot_msg_id, _ = db.save_message(user_id, persona_type, persona_type.capitalize() if persona_type != "chat" else "Nemai", response_text, persona_id=None)
        if not bot_msg_id:
             logger.error(f"Failed to save bot response to DB for user {user_id}, persona {persona_type}")
             # Continue to send response to user even if DB save fails.

        # --- Feedback View ---
        # Create the feedback view if the system is enabled and the bot message was saved.
        feedback_view = None
        if bot_msg_id and config.ENABLE_FEEDBACK_SYSTEM:
            feedback_view = FeedbackView(bot_msg_id, user_msg_id, user_id, persona_type, None)

        # --- Send Response (Pagination or Embed) ---
        # Check if the response text exceeds the Discord embed description limit.
        if len(response_text) > config.MAX_EMBED_DESCRIPTION:
            # If too long, chunk the message into pages.
            pages = chunk_message(response_text)
            if not pages:
                 logger.error("Chunking resulted in empty pages.")
                 await send_error_response(interaction, "Failed to process the response content.")
                 return False

            # Create a PaginatorView to handle page navigation.
            paginator = PaginatorView(pages, interaction.user, embed_title, color_key)
            # Send the first page and the paginator view.
            initial_message = await interaction.followup.send(embed=paginator.get_page_embed(), view=paginator, wait=True)
            paginator.message = initial_message # Store the message object in the view for later edits.

            # If feedback is enabled, send the feedback buttons as a separate message
            # because views cannot be easily combined after the initial send with pagination.
            if feedback_view and initial_message:
                 try:
                      feedback_message = await interaction.channel.send("Did you find this response helpful?", view=feedback_view)
                      feedback_view.message = feedback_message # Store message for button disabling.
                 except Exception as fb_e:
                      logger.error(f"Failed to send separate feedback message for paginated response: {fb_e}")

        else:
            # If the response fits, send it as a single embed.
            embed = discord.Embed(
                title=embed_title,
                description=response_text,
                color=config.EMBED_COLORS.get(color_key, discord.Color.blue()) # Use determined color_key
            )
            embed.set_footer(text=f"Requested by: {username}", icon_url=interaction.user.display_avatar.url)

            # Send the embed, including the feedback view if available.
            sent_message = await interaction.followup.send(embed=embed, view=feedback_view if feedback_view else discord.utils.MISSING)
            if feedback_view:
                 feedback_view.message = sent_message # Store message for button disabling.

        logger.info(f"Successfully sent '{persona_type}' response to user {user_id}")
        return True

    # --- Error Handling ---
    except genai.types.BlockedPromptException as block_err:
        # Handle errors where the AI model blocked the prompt due to safety filters.
        logger.warning(f"Request blocked for user {user_id}, type {persona_type}. Reason: {block_err}")
        await send_error_response(interaction, "I cannot process this request due to safety restrictions.")
        return False
    except Exception as e:
        # Catch any other unexpected errors during the process.
        logger.exception(f"An error occurred in handle_persona_response for {persona_type}, user {user_id}: {e}")
        await send_error_response(interaction, f"Sorry, I encountered an error while processing your request for '{persona_type}'. Please try again.")
        return False


# --- Command Implementations ---
# These functions contain the specific logic for each slash command.
# They primarily validate input and then call the appropriate handler function (like handle_persona_response).

async def chat_command_impl(interaction: Interaction, message: str):
    """Implementation for the /chat command."""
    if not message.strip(): # Prevent empty messages
        await send_error_response(interaction, "Please provide a message to chat about.")
        return
    # Call the generic handler for the 'chat' persona.
    await handle_persona_response(interaction, message, "chat", model_chat, "🤖")

async def sherlock_command_impl(interaction: Interaction, message: str):
    """Implementation for the /sherlock command."""
    if not message.strip():
        await send_error_response(interaction, "Please provide a message for Sherlock.")
        return
    # Call the generic handler for the 'sherlock' persona.
    await handle_persona_response(interaction, message, "sherlock", model_persona, "🕵️")

async def teacher_command_impl(interaction: Interaction, message: str):
    """Implementation for the /teacher command."""
    if not message.strip():
        await send_error_response(interaction, "Please provide a topic for the teacher.")
        return
    # Call the generic handler for the 'teacher' persona.
    await handle_persona_response(interaction, message, "teacher", model_chat, "👩‍🏫")

async def scientist_command_impl(interaction: Interaction, message: str):
    """Implementation for the /scientist command."""
    if not message.strip():
        await send_error_response(interaction, "Please provide a message for the scientist.")
        return
    # Call the generic handler for the 'scientist' persona.
    await handle_persona_response(interaction, message, "scientist", model_chat, "🔬")

async def search_command_impl(interaction: Interaction, query: str):
    """
    Implementation for the /search command.

    Performs a web search using DuckDuckGo, then uses the AI model to summarize the results.
    """
    # Check model availability.
    if not model_chat:
        logger.error("Cannot handle search: Chat model not available.")
        await send_error_response(interaction, "The AI model needed for summarization is currently unavailable.")
        return
    # Validate input.
    if not query.strip():
        await send_error_response(interaction, "Please provide a search query.")
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username) # Ensure user exists.

    await interaction.response.defer(thinking=True)
    logger.info(f"User {user_id} initiated web search for: '{query}'")

    search_results_text = ""
    try:
        # Perform the DuckDuckGo search in a separate thread to avoid blocking the bot.
        loop = asyncio.get_running_loop()
        results = await loop.run_in_executor(
            None, # Use default executor (ThreadPoolExecutor)
            lambda: list(DDGS().text(query, max_results=config.SEARCH_RESULT_LIMIT))
        )

        # Check if any results were found.
        if not results:
            logger.warning(f"No search results found for query: '{query}'")
            await interaction.followup.send(f"❌ No web search results found for '{query}'. Try phrasing it differently?")
            return

        # Format the search results into a string for the AI prompt.
        search_results_text += f"Search Results for Query: \"{query}\"\n\n"
        for i, result in enumerate(results):
            title = result.get('title', 'No Title')
            body = result.get('body', 'No Snippet')
            href = result.get('href', 'No URL')
            # Include angle brackets around URL for better Discord linking.
            search_results_text += f"Result {i+1}: {title}\nSnippet: {body}\nURL: <{href}>\n\n"

        logger.debug(f"Fetched {len(results)} results for query: '{query}'")

    except Exception as e:
        logger.exception(f"Error performing DuckDuckGo search for query '{query}': {e}")
        await send_error_response(interaction, "Sorry, I encountered an error while searching the web.")
        return

    # --- Summarization ---
    try:
        # Create the prompt for the AI to summarize the search results.
        summarization_prompt = (
            f"Based *only* on the following web search results provided for the query \"{query}\", "
            f"write a concise and informative summary. Synthesize the key information found in the snippets. "
            f"Do not add external knowledge. If the results seem irrelevant or contradictory, mention that briefly.\n\n"
            f"--- Search Results ---\n"
            f"{search_results_text}\n"
            f"--- End Search Results ---\n\n"
            f"Summary:"
        )

        logger.debug(f"Generating summary for search query: '{query}'")
        # Generate the summary using the chat model.
        response = await model_chat.generate_content_async(summarization_prompt)
        summary_text = response.text

        # Handle cases where the AI might return an empty summary.
        if not summary_text or summary_text.strip() == "":
             logger.warning(f"AI generated an empty summary for query: '{query}'")
             summary_text = "The AI couldn't generate a summary based on the search results."

        # Prepare embed details.
        embed_title = f"🌐 Web Search Summary: {query[:100]}{'...' if len(query)>100 else ''}" # Truncate long queries in title.
        color_key = "search"

        # Send the summary, handling pagination if necessary.
        if len(summary_text) > config.MAX_EMBED_DESCRIPTION:
            pages = chunk_message(summary_text)
            if not pages:
                 logger.error("Chunking resulted in empty pages for search summary.")
                 await send_error_response(interaction, "Failed to process the summary content.")
                 return

            view = PaginatorView(pages, interaction.user, embed_title, color_key)
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, wait=True)
        else:
            embed = discord.Embed(
                title=embed_title,
                description=summary_text,
                color=config.EMBED_COLORS.get(color_key, discord.Color.blue())
            )
            embed.set_footer(text=f"Search requested by: {username} | Results from DuckDuckGo", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=embed)

        logger.info(f"Successfully sent web search summary to user {user_id} for query: '{query}'")

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"Search summary generation blocked for query '{query}'. Reason: {block_err}")
        await send_error_response(interaction, "I cannot generate a summary for these search results due to safety restrictions.")
    except Exception as e:
        logger.exception(f"Error summarizing search results for query '{query}': {e}")
        await send_error_response(interaction, "Sorry, I encountered an error while summarizing the search results.")

async def news_command_impl(interaction: Interaction, topic: Optional[str] = None):
    """
    Implementation for the /news command.

    Fetches recent news articles using DuckDuckGo News and summarizes them using the AI model.
    """
    if not model_chat:
        logger.error("Cannot handle news command: Chat model not available.")
        await send_error_response(interaction, "The AI model needed for summarization is currently unavailable.")
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(thinking=True)

    # Use the provided topic or default to general news headlines.
    search_query = topic if topic else "Top world news headlines"
    news_article_limit = config.NEWS_ARTICLE_LIMIT
    logger.info(f"User {user_id} requested news for topic: '{search_query}' (limit: {news_article_limit})")

    news_results_text = "" # For the AI prompt
    articles_list_for_embed = [] # For displaying sources in the embed
    fetched_articles = [] # Store structured article data if needed later

    try:
        # Perform the news search.
        loop = asyncio.get_running_loop()
        results = await loop.run_in_executor(
            None,
            lambda: list(DDGS().news(
                search_query,
                region='wt-wt',      # Worldwide region
                safesearch='off',    # Safesearch setting
                timelimit='d',       # Look for news within the last day
                max_results=news_article_limit
            ))
        )

        if not results:
            logger.warning(f"No news results found for query: '{search_query}'")
            await interaction.followup.send(f"❌ No recent news found for '{search_query}'. Try a different topic or try again later.")
            return

        # Format the news results for the prompt and the embed footer.
        news_results_text += f"News Articles Found for Query: \"{search_query}\"\n\n"
        for i, result in enumerate(results):
            title = result.get('title', 'No Title')
            body = result.get('body', 'No Snippet Available')
            source = result.get('source', 'Unknown Source')
            url = result.get('url', None)
            date_str = result.get('date', None) # e.g., "2023-10-27T10:00:00"

            # Store structured data.
            fetched_articles.append({
                'title': title, 'source': source, 'url': url, 'date': date_str, 'body': body
            })

            # Build the text for the AI prompt.
            article_text = f"Article {i+1}:\n"
            article_text += f"Title: {title}\n"
            article_text += f"Source: {source}\n"
            if date_str:
                article_text += f"Date: {date_str}\n"
            article_text += f"Snippet: {body}\n"
            if url:
                article_text += f"URL: {url}\n\n"
            else:
                article_text += "\n"
            news_results_text += article_text

            # Build the markdown list for the embed description (if space allows).
            if url:
                # Format as "[Source: Title](URL)"
                articles_list_for_embed.append(f"• [{source}: {title}]({url})")
            else:
                articles_list_for_embed.append(f"• {source}: {title} (No URL)")

        logger.debug(f"Fetched {len(results)} news articles for query: '{search_query}'")

    except Exception as e:
        logger.exception(f"Error performing DuckDuckGo news search for query '{search_query}': {e}")
        await send_error_response(interaction, "Sorry, I encountered an error while searching for news.")
        return

    # --- Summarization ---
    try:
        # Create the prompt for summarizing the news articles.
        summarization_prompt = (
            f"You are tasked with summarizing the following news articles related to the query \"{search_query}\".\n"
            f"Based *only* on the text provided below, write a concise and neutral news summary. Synthesize the main points from the different articles.\n"
            f"Do not add external information or opinions. Start the summary directly, without introductory phrases like 'Here is a summary...'.\n\n"
            f"--- News Articles ---\n"
            f"{news_results_text}\n"
            f"--- End News Articles ---\n\n"
            f"Concise News Summary:"
        )

        logger.debug(f"Generating summary for news query: '{search_query}'")
        response = await model_chat.generate_content_async(summarization_prompt)
        summary_text = response.text.strip()

        if not summary_text:
             logger.warning(f"AI generated an empty news summary for query: '{search_query}'")
             summary_text = "The AI couldn't generate a summary based on the news articles found."

        # Attempt to add source links to the embed description if they fit.
        if articles_list_for_embed:
            # Limit to first 5 sources to avoid excessive length.
            sources_markdown = "\n\n**Sources:**\n" + "\n".join(articles_list_for_embed[:5])
            # Check if adding sources exceeds the embed description limit.
            if len(summary_text) + len(sources_markdown) <= config.MAX_EMBED_DESCRIPTION:
                 summary_text += sources_markdown
            else:
                 logger.info("Skipping source links in description due to length constraints.")

        # Prepare embed details.
        embed_title_topic = topic if topic else "Top Headlines"
        embed_title = f"📰 News Summary: {embed_title_topic[:100]}{'...' if len(embed_title_topic)>100 else ''}"
        color_key = "news"

        # Send the summary, handling pagination.
        if len(summary_text) > config.MAX_EMBED_DESCRIPTION:
            # If paginating, use the summary *without* sources for the pages,
            # as sources might have been truncated or skipped.
            summary_text_only_for_pages = response.text.strip()
            pages = chunk_message(summary_text_only_for_pages)
            if not pages:
                 logger.error("Chunking resulted in empty pages for news summary.")
                 await send_error_response(interaction, "Failed to process the summary content.")
                 return

            # Create a temporary embed for the first page to pass to the PaginatorView constructor.
            # This isn't strictly necessary but ensures the PaginatorView is initialized correctly.
            first_page_embed = PaginatorView([pages[0]], interaction.user, embed_title, color_key).get_page_embed()

            # Create the actual paginator view.
            view = PaginatorView(pages, interaction.user, embed_title, color_key)
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, wait=True)
        else:
            # Send as a single embed.
            embed = discord.Embed(
                title=embed_title,
                description=summary_text, # Contains sources if they fit.
                color=config.EMBED_COLORS.get(color_key, discord.Color.dark_blue())
            )
            embed.set_footer(text=f"News requested by: {username} | Articles via DuckDuckGo News", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=embed)

        logger.info(f"Successfully sent news summary to user {user_id} for query: '{search_query}'")

    except genai.types.BlockedPromptException as block_err:
        # Handle blocked prompts, but try to send the source links if available.
        logger.warning(f"News summary generation blocked for query '{search_query}'. Reason: {block_err}")
        if articles_list_for_embed:
            # Create an embed containing only the source links.
            error_summary = "I couldn't generate a summary due to safety restrictions, but here are the links found:\n\n" + "\n".join(articles_list_for_embed)
            error_embed = discord.Embed(
                 title=f"📰 News Links: {embed_title_topic[:100]}{'...' if len(embed_title_topic)>100 else ''}",
                 description=error_summary[:config.MAX_EMBED_DESCRIPTION], # Truncate if needed
                 color=config.EMBED_COLORS.get("warning", discord.Color.orange())
            )
            error_embed.set_footer(text=f"News requested by: {username} | Articles via DuckDuckGo News", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=error_embed)
        else:
            # If no links either, send a standard error.
            await send_error_response(interaction, "I cannot generate a summary for these news articles due to safety restrictions, and no links were found.")

    except Exception as e:
        logger.exception(f"Error summarizing news results for query '{search_query}': {e}")
        await send_error_response(interaction, "Sorry, I encountered an error while summarizing the news.")


async def analyze_file_command_impl(interaction: Interaction, file: Attachment, prompt: Optional[str] = None):
    """
    Implementation for the /analyze_file command.

    Analyzes the content of an uploaded file (image or text-based) using the appropriate AI model.
    """
    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    file_type = file.content_type
    is_image = file_type in config.ALLOWED_IMAGE_TYPES
    is_text_based = file_type in config.ALLOWED_TEXT_TYPES

    # --- Validation ---
    # Check if the file type is supported.
    if not is_image and not is_text_based:
        allowed_types_str = ", ".join(sorted(list(set([t.split('/')[-1] for t in config.ALL_ALLOWED_FILE_TYPES])))) or "None configured"
        await send_error_response(interaction, f"Unsupported file type ('{file_type}'). Please upload one of the allowed types: {allowed_types_str}.")
        return

    # Check if the file size exceeds the configured limit.
    if file.size > config.MAX_FILE_SIZE_BYTES:
         await send_error_response(interaction, f"File is too large ({file.size / (1024*1024):.2f} MB). Maximum allowed size is {config.MAX_FILE_SIZE_BYTES / (1024*1024):.2f} MB.")
         return

    await interaction.response.defer(thinking=True)

    try:
        # Read the file content into bytes.
        file_bytes = await file.read()
        file_content_str = None # Will hold extracted text for text-based files.

        # --- Processing based on file type ---
        if is_image:
            # Use the vision model for images.
            if not model_vision:
                logger.error("Cannot handle image analysis: Vision model not available.")
                await send_error_response(interaction, "The AI model for image analysis is currently unavailable.")
                return

            # Prepare the image data part for the Gemini API.
            img_part = {"mime_type": file_type, "data": file_bytes}
            # Combine the text prompt (or default) and the image data.
            prompt_content = [prompt if prompt else "Describe this image in detail.", img_part]
            model_to_use = model_vision
            embed_title = f"🖼️ Image Analysis: {file.filename}"
            color_key = "image"
            log_msg_start = f"Generating image analysis for user {user_id} on file '{file.filename}'. Prompt: {prompt if prompt else 'Describe image'}"
            log_msg_success = f"Successfully sent image analysis response for '{file.filename}' to user {user_id}."

        elif is_text_based:
            # Use the chat model for text files.
            if not model_chat:
                logger.error("Cannot handle text file analysis: Chat model not available.")
                await send_error_response(interaction, "The AI model for text analysis is currently unavailable.")
                return

            # Extract text content based on the file type.
            try:
                if file_type == "application/pdf":
                    # Use PyPDF2 to extract text from PDF.
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                    file_content_str = "".join([page.extract_text() for page in pdf_reader.pages])
                elif file_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                    # Use python-docx to extract text from DOC/DOCX.
                    doc = docx.Document(io.BytesIO(file_bytes))
                    file_content_str = "\n".join([para.text for para in doc.paragraphs])
                else:
                     # For other text types, try decoding as UTF-8, fallback to latin-1.
                     try:
                         file_content_str = file_bytes.decode('utf-8')
                     except UnicodeDecodeError:
                         logger.warning(f"Could not decode file '{file.filename}' as UTF-8. Trying latin-1.")
                         file_content_str = file_bytes.decode('latin-1') # Common fallback.

            except Exception as parse_err:
                 # Handle errors during text extraction.
                 logger.error(f"Failed to parse file '{file.filename}' (type: {file_type}): {parse_err}", exc_info=True)
                 await send_error_response(interaction, f"Could not extract text from '{file.filename}'. The file might be corrupted, password-protected, or an unsupported variation of {file_type}.")
                 return

            # Check if text extraction yielded any content.
            if not file_content_str or not file_content_str.strip():
                 logger.warning(f"Extracted empty text content from file '{file.filename}' (type: {file_type}).")
                 await send_error_response(interaction, f"Could not find any readable text content in '{file.filename}'.")
                 return

            # Truncate very long text content to avoid exceeding model limits.
            if len(file_content_str) > config.MAX_FILE_CONTENT_CHARS:
                file_content_str = file_content_str[:config.MAX_FILE_CONTENT_CHARS]
                logger.warning(f"File content for '{file.filename}' truncated to {config.MAX_FILE_CONTENT_CHARS} characters for analysis.")

            # Construct the prompt for the AI, including the extracted text.
            if prompt:
                # If user provided a specific prompt.
                prompt_content = (
                    f"Analyze the following file content based on the user's request.\n\n"
                    f"--- FILE CONTENT ({file.filename}) START ---\n"
                    f"{file_content_str}\n"
                    f"--- FILE CONTENT ({file.filename}) END ---\n\n"
                    f"User's Request: {prompt}\n\n"
                    f"Analysis:"
                )
            else:
                # Default prompt: summarize and analyze.
                prompt_content = (
                    f"Provide a concise summary and analysis of the following file content.\n\n"
                    f"--- FILE CONTENT ({file.filename}) START ---\n"
                    f"{file_content_str}\n"
                    f"--- FILE CONTENT ({file.filename}) END ---\n\n"
                    f"Summary and Analysis:"
                )

            model_to_use = model_chat
            embed_title = f"📄 File Analysis: {file.filename}"
            color_key = "info"
            log_msg_start = f"Generating text file analysis for user {user_id} on file '{file.filename}'. Prompt: {prompt if prompt else 'Summarize file'}"
            log_msg_success = f"Successfully sent text file analysis response for '{file.filename}' to user {user_id}."

        else:
            # This case should not be reachable due to initial validation.
            logger.error(f"File type '{file_type}' passed initial check but failed secondary logic for user {user_id}.")
            await send_error_response(interaction, "An internal error occurred determining the file type for analysis.")
            return

        # --- AI Generation and Response ---
        logger.debug(log_msg_start)
        # Generate the analysis using the selected model and prompt.
        response = await model_to_use.generate_content_async(prompt_content)
        response_text = response.text

        # Save a record of the analysis request/response (optional, but good for tracking).
        # Save only a snippet of the response to avoid large DB entries.
        bot_msg_id, _ = db.save_message(
            user_id,
            "doc_assist", # Use a relevant conversation type
            "Nemai",
            f"Analysis for {file.filename}:\n{response_text[:500]}...", # Store truncated response
            persona_id=None
        )

        # Send the response, handling pagination.
        if len(response_text) > config.MAX_EMBED_DESCRIPTION:
            pages = chunk_message(response_text)
            if not pages:
                 logger.error(f"Chunking resulted in empty pages for file analysis: {file.filename}")
                 await send_error_response(interaction, "Failed to process the analysis content.")
                 return

            view = PaginatorView(pages, interaction.user, embed_title, color_key)
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, wait=True)
        else:
            embed = discord.Embed(
                title=embed_title,
                description=response_text,
                color=config.EMBED_COLORS.get(color_key, discord.Color.light_grey())
            )
            embed.set_footer(text=f"Requested by: {username}", icon_url=interaction.user.display_avatar.url)

            # If it was an image analysis, try to display the image in the embed.
            if is_image:
                try:
                     # Use the file's URL provided by Discord.
                     embed.set_image(url=file.url)
                except Exception as img_err: # Catch potential errors setting the URL
                     logger.warning(f"Could not set image URL for embed: {file.url}. Error: {img_err}")

            await interaction.followup.send(embed=embed)

        logger.info(log_msg_success)

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"File analysis blocked for user {user_id} on file '{file.filename}'. Reason: {block_err}")
        await send_error_response(interaction, "I cannot process this request due to safety restrictions on the file content or prompt.")
    except discord.HTTPException as discord_err:
         # Handle errors reading the file from Discord.
         logger.error(f"Discord HTTP error processing file '{file.filename}' for user {user_id}: {discord_err}")
         await send_error_response(interaction, f"Failed to read or process the uploaded file '{file.filename}'. Please try again.")
    except Exception as e:
        logger.exception(f"An error occurred during file analysis for user {user_id} on file '{file.filename}': {e}")
        await send_error_response(interaction, f"Sorry, I encountered an error while analyzing the file '{file.filename}'.")


async def factcheck_command_impl(interaction: Interaction, statement: str):
    """
    Implementation for the /factcheck command.

    Uses web search results to evaluate the truthfulness of a given statement.
    """
    if not model_chat:
        logger.error("Cannot handle fact-check: Chat model not available.")
        await send_error_response(interaction, "The AI model needed for analysis is currently unavailable.")
        return
    if not statement.strip():
        await send_error_response(interaction, "Please provide a statement to fact-check.")
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(thinking=True)
    logger.info(f"User {user_id} initiated fact-check for: '{statement}'")

    # --- Web Search ---
    search_results_text = ""
    try:
        # Search the web for the statement itself to find relevant context/evidence.
        loop = asyncio.get_running_loop()
        results = await loop.run_in_executor(
            None,
            lambda: list(DDGS().text(statement, max_results=config.SEARCH_RESULT_LIMIT))
        )

        if not results:
            logger.warning(f"No search results found for fact-check statement: '{statement}'")
            # If no results, we cannot perform the fact-check.
            await interaction.followup.send(f"❌ No web search results found to evaluate the statement: '{statement}'. Cannot perform fact-check.")
            return

        # Format results for the AI prompt.
        search_results_text += f"Search Results for Statement: \"{statement}\"\n\n"
        for i, result in enumerate(results):
            title = result.get('title', 'No Title')
            body = result.get('body', 'No Snippet')
            href = result.get('href', 'No URL')
            search_results_text += f"Result {i+1}: {title}\nSnippet: {body}\nURL: <{href}>\n\n"

        logger.debug(f"Fetched {len(results)} results for fact-check: '{statement}'")

    except Exception as e:
        logger.exception(f"Error performing DuckDuckGo search for fact-check statement '{statement}': {e}")
        await send_error_response(interaction, "Sorry, I encountered an error while searching the web for information.")
        return

    # --- AI Analysis ---
    try:
        # Construct the prompt for the AI to analyze the statement based *only* on the search results.
        factcheck_prompt = (
            f"Analyze the following statement based *only* on the provided web search results. Do not use external knowledge.\n\n"
            f"Statement to Evaluate: \"{statement}\"\n\n"
            f"--- Web Search Results ---\n"
            f"{search_results_text}\n"
            f"--- End Web Search Results ---\n\n"
            f"Instructions:\n"
            f"1. Evaluate the truthfulness of the statement based *strictly* on the information within the search results.\n"
            f"2. Classify the statement's veracity using ONE of the following labels: **True**, **False**, **Partially True/Nuanced**, **Misleading**, **Unverifiable** (if results are irrelevant or insufficient).\n"
            f"3. Provide a concise explanation for your classification, directly referencing snippets or titles from the search results (e.g., 'Result 3 suggests...', 'Results 1 and 5 contradict...').\n"
            f"4. If classified as Partially True/Nuanced or Misleading, explain the nuance or how it's misleading according to the results.\n"
            f"5. If Unverifiable, explain why the results don't allow for a conclusion.\n"
            f"6. Be objective and present the findings clearly.\n\n"
            f"Analysis:\n"
            f"**Classification:** [Your Classification Label Here]\n" # Guide the AI's output format
            f"**Explanation:** [Your explanation based on the results]"
        )

        logger.debug(f"Generating fact-check analysis for statement: '{statement}'")
        response = await model_chat.generate_content_async(factcheck_prompt)
        analysis_text = response.text

        if not analysis_text or analysis_text.strip() == "":
             logger.warning(f"AI generated an empty fact-check analysis for statement: '{statement}'")
             analysis_text = "The AI could not generate an analysis based on the search results."

        # Prepare embed details.
        embed_title = f"✔️ Fact Check Analysis"
        color_key = "factcheck"

        # Combine the original statement and the AI's analysis for the embed description.
        full_description = f"**Statement:** {statement}\n\n{analysis_text}"

        # Send response, handling pagination.
        if len(full_description) > config.MAX_EMBED_DESCRIPTION:
            pages = chunk_message(full_description)
            if not pages:
                 logger.error("Chunking resulted in empty pages for fact-check analysis.")
                 await send_error_response(interaction, "Failed to process the analysis content.")
                 return

            view = PaginatorView(pages, interaction.user, embed_title, color_key)
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, wait=True)
        else:
            embed = discord.Embed(
                title=embed_title,
                description=full_description,
                color=config.EMBED_COLORS.get(color_key, discord.Color.dark_green())
            )
            embed.set_footer(text=f"Fact-check requested by: {username} | Analysis based on DuckDuckGo results", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=embed)

        logger.info(f"Successfully sent fact-check analysis to user {user_id} for statement: '{statement}'")

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"Fact-check analysis generation blocked for statement '{statement}'. Reason: {block_err}")
        await send_error_response(interaction, "I cannot generate an analysis for this statement due to safety restrictions.")
    except Exception as e:
        logger.exception(f"Error generating fact-check analysis for statement '{statement}': {e}")
        await send_error_response(interaction, "Sorry, I encountered an error while analyzing the statement.")


async def history_search_command_impl(interaction: Interaction, query: str, limit: int):
    """
    Implementation for the /search_history command.

    Searches the user's past conversation history using semantic similarity.
    """
    if embedding_model is None:
        logger.error("Cannot search history: Embedding model not available.")
        await send_error_response(interaction, "The text analysis component needed for history search is currently unavailable.", ephemeral=True)
        return
    if not query.strip():
        await send_error_response(interaction, "Please provide a search query for your history.", ephemeral=True)
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    # Respond ephemerally as history is private.
    await interaction.response.defer(ephemeral=True, thinking=True)
    logger.info(f"User {user_id} initiated history search for: '{query}' (limit: {limit})")

    # --- Embedding and Search ---
    try:
        # Generate embedding for the user's search query.
        query_embedding = embedding_model.encode(query, convert_to_tensor=False).astype(np.float32)
    except Exception as e:
        logger.error(f"Failed to create embedding for history search query '{query}' by user {user_id}: {e}")
        await send_error_response(interaction, "Failed to process your search query. Please try again.", ephemeral=True)
        return

    try:
        # Perform the semantic search in the database.
        history_results = db.search_user_history(user_id, query_embedding, limit)

        if not history_results:
            await interaction.followup.send(f"❌ No relevant messages found in your history matching '{query}'.", ephemeral=True)
            return

        # --- Formatting Results ---
        formatted_results = []
        max_snippet_len = 150 # Max length for message snippets in the results.
        for role, content, timestamp_str, conv_type, similarity in history_results:
            # Format timestamp nicely using Discord's relative time format.
            try:
                dt_obj = datetime.datetime.fromisoformat(timestamp_str)
                # <t:unix_timestamp:R> format for relative time (e.g., "5 minutes ago").
                ts_formatted = f"<t:{int(dt_obj.timestamp())}:R>"
            except ValueError:
                ts_formatted = "Unknown Time" # Fallback if timestamp parsing fails.

            # Create a snippet of the message content.
            content_snippet = content.replace('\n', ' ')[:max_snippet_len]
            if len(content) > max_snippet_len:
                content_snippet += "..."

            # Add conversation type and similarity score.
            type_display = f"({conv_type})" if conv_type else "(Unknown Type)"
            similarity_display = f"(Similarity: {similarity:.2f})"

            formatted_results.append(f"**[{ts_formatted}]** {type_display} **{role}:** {content_snippet} {similarity_display}")

        # Combine formatted results into a single string.
        result_text = "\n\n".join(formatted_results)
        embed_title = f"📜 History Search Results for: \"{query[:100]}{'...' if len(query)>100 else ''}\""
        color_key = "history_search"

        # Send results, handling pagination if needed (ephemerally).
        if len(result_text) > config.MAX_EMBED_DESCRIPTION:
            # Adjust chunk size for ephemeral messages if needed (though MAX_PAGE_SIZE should be fine).
            pages = chunk_message(result_text, chunk_size=config.MAX_PAGE_SIZE - 100) # Slightly smaller chunk for safety.
            if not pages:
                 logger.error(f"Chunking resulted in empty pages for history search: {query}")
                 await send_error_response(interaction, "Failed to process the history search results.", ephemeral=True)
                 return

            view = PaginatorView(pages, interaction.user, embed_title, color_key, timeout=config.PAGINATOR_TIMEOUT)
            # Send the paginator ephemerally.
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, ephemeral=True, wait=True)

        else:
            embed = discord.Embed(
                title=embed_title,
                description=result_text,
                color=config.EMBED_COLORS.get(color_key, discord.Color.dark_gold())
            )
            embed.set_footer(text=f"Found {len(history_results)} results | Search by: {username}", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=embed, ephemeral=True)

        logger.info(f"Successfully sent history search results to user {user_id} for query: '{query}'")

    except Exception as e:
        logger.exception(f"Error during history search for user {user_id}, query '{query}': {e}")
        await send_error_response(interaction, "Sorry, an error occurred while searching your history.", ephemeral=True)


async def explain_like_im_impl(interaction: Interaction, topic: str, audience: str):
    """
    Implementation for the /explain_like_im command.

    Explains a topic using web search results, tailored to a specific audience description.
    """
    if not model_chat:
        logger.error("Cannot handle explain_like_im: Chat model not available.")
        await send_error_response(interaction, "The AI model needed for explanation is currently unavailable.")
        return
    if not topic.strip():
        await send_error_response(interaction, "Please provide a topic to explain.")
        return
    if not audience.strip():
        await send_error_response(interaction, "Please specify the audience (e.g., 'a 5-year-old', 'a physics professor').")
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(thinking=True)
    logger.info(f"User {user_id} requested explanation of '{topic}' for audience '{audience}'")

    # --- Web Search for Context ---
    search_results_text = ""
    try:
        # Search the web for the topic to gather information.
        loop = asyncio.get_running_loop()
        results = await loop.run_in_executor(
            None,
            lambda: list(DDGS().text(topic, max_results=config.SEARCH_RESULT_LIMIT))
        )

        if not results:
            logger.warning(f"No search results found for ELI5 topic: '{topic}'")
            await interaction.followup.send(f"❌ No web search results found for '{topic}'. Cannot generate explanation.")
            return

        # Format results for the AI prompt.
        search_results_text += f"Search Results for Topic: \"{topic}\"\n\n"
        for i, result in enumerate(results):
            title = result.get('title', 'No Title')
            body = result.get('body', 'No Snippet')
            href = result.get('href', 'No URL')
            search_results_text += f"Result {i+1}: {title}\nSnippet: {body}\nURL: <{href}>\n\n"

        logger.debug(f"Fetched {len(results)} results for ELI5 topic: '{topic}'")

    except Exception as e:
        logger.exception(f"Error performing DuckDuckGo search for ELI5 topic '{topic}': {e}")
        await send_error_response(interaction, "Sorry, I encountered an error while searching the web for information.")
        return

    # --- AI Explanation Generation ---
    try:
        # Construct the prompt, instructing the AI to explain based on results and tailor to the audience.
        explanation_prompt = (
            f"Based *only* on the following web search results provided for the topic \"{topic}\", "
            f"explain this topic as if you were speaking directly to the following audience: **{audience}**. "
            f"Adopt their likely vocabulary, level of detail, perspective, and tone. "
            f"Synthesize the key information from the snippets relevant to this audience. "
            f"Do not add external knowledge. Start the explanation directly.\n\n"
            f"--- Search Results ---\n"
            f"{search_results_text}\n"
            f"--- End Search Results ---\n\n"
            f"Explanation for '{audience}':"
        )

        logger.debug(f"Generating ELI5 explanation for topic '{topic}', audience '{audience}'")
        response = await model_chat.generate_content_async(explanation_prompt)
        explanation_text = response.text.strip()

        if not explanation_text:
            logger.warning(f"AI generated an empty explanation for topic: '{topic}', audience: '{audience}'")
            explanation_text = "The AI couldn't generate an explanation based on the search results for this audience."

        # Prepare embed details.
        embed_title = f"🧠 Explaining '{topic[:100]}{'...' if len(topic)>100 else ''}' for {audience[:100]}{'...' if len(audience)>100 else ''}"
        color_key = "definition" # Often similar to definitions.

        # Send response, handling pagination.
        if len(explanation_text) > config.MAX_EMBED_DESCRIPTION:
            pages = chunk_message(explanation_text)
            if not pages:
                 logger.error("Chunking resulted in empty pages for ELI5 explanation.")
                 await send_error_response(interaction, "Failed to process the explanation content.")
                 return

            view = PaginatorView(pages, interaction.user, embed_title, color_key)
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, wait=True)
        else:
            embed = discord.Embed(
                title=embed_title,
                description=explanation_text,
                color=config.EMBED_COLORS.get(color_key, discord.Color.purple())
            )
            embed.set_footer(text=f"Explanation requested by: {username}", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=embed)

        logger.info(f"Successfully sent ELI5 explanation to user {user_id} for topic '{topic}', audience '{audience}'")

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"ELI5 explanation generation blocked for topic '{topic}'. Reason: {block_err}")
        await send_error_response(interaction, "I cannot generate an explanation for this topic/audience due to safety restrictions.")
    except Exception as e:
        logger.exception(f"Error generating ELI5 explanation for topic '{topic}': {e}")
        await send_error_response(interaction, "Sorry, I encountered an error while generating the explanation.")


# --- Persona Management Command Implementations ---

async def persona_create_impl(interaction: Interaction, name: str, description: str):
    """Implementation for the /persona create command."""
    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    name = name.strip()
    description = description.strip()

    # Validate name and description lengths based on config.
    if not (config.PERSONA_NAME_MIN_LEN <= len(name) <= config.PERSONA_NAME_MAX_LEN):
        await send_error_response(interaction, f"Persona name must be between {config.PERSONA_NAME_MIN_LEN} and {config.PERSONA_NAME_MAX_LEN} characters.")
        return
    if not (config.PERSONA_DESC_MIN_LEN <= len(description) <= config.PERSONA_DESC_MAX_LEN):
        await send_error_response(interaction, f"Persona description must be between {config.PERSONA_DESC_MIN_LEN} and {config.PERSONA_DESC_MAX_LEN} characters.")
        return

    await interaction.response.defer(ephemeral=True) # Respond ephemerally as it's user-specific management.
    # Attempt to add the persona to the database.
    success, result = db.add_persona(user_id, name, description)

    if success:
        slot_id = result # `add_persona` returns the slot_id on success.
        count = db.get_persona_count(user_id)
        await interaction.followup.send(
            f"✅ Persona '{name}' created successfully in **Slot {slot_id}**! ({count}/{db.MAX_PERSONAS} slots used)\n"
            f"Use `/persona activate slot_id:{slot_id}` to start chatting with it.",
            ephemeral=True
        )
    else:
        # `add_persona` returns an error message string on failure.
        await send_error_response(interaction, f"Failed to create persona: {result}")


async def persona_delete_impl(interaction: Interaction, slot_id: app_commands.Range[int, 1, config.MAX_PERSONAS]):
    """Implementation for the /persona delete command."""
    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(ephemeral=True)
    # Attempt to delete the persona from the database.
    success, result_message = db.delete_persona(user_id, slot_id)

    if success:
        await interaction.followup.send(f"✅ {result_message}", ephemeral=True)
    else:
        await send_error_response(interaction, result_message)


async def persona_list_impl(interaction: Interaction):
    """Implementation for the /persona list command."""
    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(ephemeral=True)
    # Get all personas for the user from the database.
    personas = db.get_personas(user_id)

    if not personas:
        await interaction.followup.send("You haven't created any personas yet. Use `/persona create` to make one!", ephemeral=True)
        return

    # Create an embed to display the persona list.
    embed = discord.Embed(title=f"{username}'s Personas", color=config.EMBED_COLORS.get("persona_list", discord.Color.purple()))
    embed.description = f"You have {len(personas)}/{db.MAX_PERSONAS} personas."
    for p_slot_id, name, created_at_str, is_active in personas:
        status_emoji = "✅ (Active)" if is_active else "❌ (Inactive)"
        # Format the creation date nicely.
        try:
             created_date = datetime.datetime.fromisoformat(created_at_str).strftime('%Y-%m-%d')
        except:
             created_date = "Unknown date" # Fallback

        embed.add_field(
            name=f"Slot {p_slot_id}: {name}",
            value=f"Status: {status_emoji}\nCreated: {created_date}",
            inline=False # Display each persona in its own block.
        )
    await interaction.followup.send(embed=embed, ephemeral=True)


async def persona_info_impl(interaction: Interaction, slot_id: app_commands.Range[int, 1, config.MAX_PERSONAS]):
    """Implementation for the /persona info command."""
    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(ephemeral=True)
    # Get detailed information for the specified persona slot.
    persona_details = db.get_persona_details_by_slot(user_id, slot_id)

    if not persona_details:
        await send_error_response(interaction, f"Persona in slot {slot_id} not found.")
        return

    # Unpack the details.
    _internal_id, p_slot_id, p_name, p_prompt, created_at_str, is_active = persona_details
    status_emoji = "✅ (Active)" if is_active else "❌ (Inactive)"
    # Format timestamp with time.
    try:
        created_date = datetime.datetime.fromisoformat(created_at_str).strftime('%Y-%m-%d %H:%M UTC')
    except:
        created_date = "Unknown"

    # Create embed to display details.
    embed = discord.Embed(
        title=f"Persona Info: {p_name} (Slot {p_slot_id})",
        description=f"**Status:** {status_emoji}\n**Created:** {created_date}\n\n**Instructions/Prompt:**",
        # Use different colors based on active status.
        color=config.EMBED_COLORS.get("persona_custom") if is_active else config.EMBED_COLORS.get("info")
    )

    # Display the persona's prompt/instructions (truncated if too long).
    # Use a code block for better formatting.
    prompt_display = p_prompt[:1020] + ("..." if len(p_prompt) > 1020 else "") # Max field value length is 1024.
    embed.add_field(name="-", value=f"```{prompt_display}```", inline=False) # Using "-" as name for visual spacing.
    await interaction.followup.send(embed=embed, ephemeral=True)


async def persona_activate_impl(interaction: Interaction, slot_id: app_commands.Range[int, 1, config.MAX_PERSONAS]):
    """Implementation for the /persona activate command."""
    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(ephemeral=True)
    # Attempt to set the specified persona as active in the database.
    success, result_message = db.set_active_persona(user_id, slot_id)

    if success:
        await interaction.followup.send(f"✅ {result_message}", ephemeral=True)
    else:
        await send_error_response(interaction, result_message)


async def persona_chat_impl(interaction: Interaction, message: str):
    """Implementation for the /persona chat command."""
    # --- Initial Checks ---
    if not model_persona:
         logger.error("Cannot handle persona chat: Persona AI model not available.")
         await send_error_response(interaction, "The AI model for personas is currently unavailable.")
         return
    if embedding_model is None:
        logger.error(f"Cannot handle persona chat: Embedding model not available.")
        await send_error_response(interaction, "The text analysis component is currently unavailable.")
        return # Changed from False to None as it doesn't return bool
    if not message.strip():
        await send_error_response(interaction, "Please provide a message for your persona.")
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(thinking=True)
    # Get the user's currently active persona from the database.
    active_persona = db.get_active_persona(user_id)
    if not active_persona:
        # Provide helpful messages if no persona is active.
        if db.get_persona_count(user_id) > 0:
             await send_error_response(interaction, "You have personas, but none are active. Use `/persona activate slot_id:[slot_number]` first.")
        else:
             await send_error_response(interaction, "You don't have an active persona. Use `/persona create` and `/persona activate` first.")
        return

    # Unpack active persona details.
    internal_p_id, p_slot_id, p_name, p_prompt = active_persona

    # --- Message Saving and Context ---
    # Save user message, linking it to the specific persona ID.
    user_msg_id, user_msg_embedding = db.save_message(user_id, "persona", "User", message, persona_id=internal_p_id)
    if not user_msg_id:
        logger.error(f"Failed to save user message to DB for persona chat, user {user_id}, persona slot {p_slot_id}")
        await send_error_response(interaction, "Failed to record your message. Please try again.")
        return

    try:
        # Get relevant history specifically for the 'persona' conversation type.
        relevant_history_list = db.get_relevant_history(user_id, message, user_msg_embedding, "persona")
        if not relevant_history_list:
             relevant_history_list = db.get_conversation_history(user_id, "persona")
        history_context = "\n".join(relevant_history_list)

        # Build feedback context specific to this persona.
        feedback_context = build_feedback_prompt_section(user_id, "persona", internal_p_id, user_msg_embedding)

        # --- Prompt Generation ---
        # Construct the prompt, including persona instructions, history, and feedback.
        persona_prompt = (
            f"You are currently acting as the persona '{p_name}'. Follow these instructions precisely:\n"
            f"--- PERSONA INSTRUCTIONS START ---\n{p_prompt}\n--- PERSONA INSTRUCTIONS END ---\n\n"
            f"Maintain this persona throughout your response. Use the following conversation history for context:\n"
            f"Relevant Conversation History (if any):\n{history_context}\n"
            f"User's username is {username}\n" # Include username for potential personalization.
            f"{feedback_context}\n\n" # Include feedback guidance.
            f"Now, respond to the user's latest message as '{p_name}': {message}"
        )

        # --- AI Response and Saving ---
        logger.debug(f"Generating persona '{p_name}' (Slot {p_slot_id}, Internal ID: {internal_p_id}) response for user {user_id}.")
        response = await model_persona.generate_content_async(persona_prompt)
        response_text = response.text

        # Save the bot's response, linking it to the persona ID.
        bot_msg_id, _ = db.save_message(user_id, "persona", p_name, response_text, persona_id=internal_p_id)
        if not bot_msg_id:
            logger.error(f"Failed to save persona '{p_name}' response to DB for user {user_id}")
            # Continue anyway.

        # --- Feedback and Sending ---
        feedback_view = None
        if bot_msg_id and config.ENABLE_FEEDBACK_SYSTEM:
             # Create feedback view specific to this persona interaction.
             feedback_view = FeedbackView(bot_msg_id, user_msg_id, user_id, "persona", internal_p_id)

        embed_title = f"{p_name}'s Response (Slot {p_slot_id}) 🎭"
        color_key = "persona_custom"

        # Send response, handling pagination and feedback view attachment.
        if len(response_text) > config.MAX_EMBED_DESCRIPTION:
            pages = chunk_message(response_text)
            if not pages:
                 logger.error("Chunking resulted in empty pages for persona response.")
                 await send_error_response(interaction, "Failed to process the persona response content.")
                 return

            paginator = PaginatorView(pages, interaction.user, embed_title, color_key)
            initial_message = await interaction.followup.send(embed=paginator.get_page_embed(), view=paginator, wait=True)
            paginator.message = initial_message

            # Send feedback separately for paginated messages.
            if feedback_view and initial_message:
                 try:
                      feedback_message = await interaction.channel.send("Did you find this response helpful?", view=feedback_view)
                      feedback_view.message = feedback_message
                 except Exception as fb_e:
                      logger.error(f"Failed to send separate feedback message for paginated persona response: {fb_e}")

        else:
            embed = discord.Embed(
                title=embed_title,
                description=response_text,
                color=config.EMBED_COLORS.get(color_key, discord.Color.blue())
            )
            embed.set_footer(text=f"Interacting as {p_name} (Slot {p_slot_id}) | Requested by: {username}", icon_url=interaction.user.display_avatar.url)
            sent_message = await interaction.followup.send(embed=embed, view=feedback_view if feedback_view else discord.utils.MISSING)
            if feedback_view:
                 feedback_view.message = sent_message

        logger.info(f"Successfully sent persona '{p_name}' response to user {user_id}")

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"Persona chat blocked for user {user_id}, persona {p_name}. Reason: {block_err}")
        await send_error_response(interaction, f"I cannot process this request for persona '{p_name}' due to safety restrictions.")
    except Exception as e:
        logger.exception(f"Error during persona chat for '{p_name}' (Slot {p_slot_id}), user {user_id}: {e}")
        await send_error_response(interaction, f"Sorry, an error occurred while chatting with '{p_name}'.")


async def persona_debate_impl(interaction: Interaction, slot_id_1: app_commands.Range[int, 1, config.MAX_PERSONAS], slot_id_2: app_commands.Range[int, 1, config.MAX_PERSONAS], topic: str):
    """Implementation for the /persona debate command."""
    # --- Initial Checks ---
    if not model_debate:
        logger.error("Cannot handle persona debate: Debate AI model not available.")
        await send_error_response(interaction, "The AI model for debates is currently unavailable.")
        return
    if embedding_model is None:
        logger.error("Cannot handle persona debate: Embedding model not available.")
        await send_error_response(interaction, "The text analysis component is currently unavailable.")
        return # Changed from False to None
    if not topic.strip():
        await send_error_response(interaction, "Please provide a topic for the debate.")
        return
    if slot_id_1 == slot_id_2: # Personas must be different.
        await send_error_response(interaction, "Personas must be in different slots to debate.")
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(thinking=True)

    # --- Get Persona Details ---
    persona1_details = db.get_persona_details_by_slot(user_id, slot_id_1)
    persona2_details = db.get_persona_details_by_slot(user_id, slot_id_2)

    if not persona1_details:
        await send_error_response(interaction, f"Persona in slot {slot_id_1} not found.")
        return
    if not persona2_details:
        await send_error_response(interaction, f"Persona in slot {slot_id_2} not found.")
        return

    # Unpack details for both personas.
    p1_id, _, p1_name, p1_prompt, _, _ = persona1_details
    p2_id, _, p2_name, p2_prompt, _, _ = persona2_details

    # --- Feedback Context (Optional) ---
    # Get embedding for the debate topic to find relevant feedback.
    debate_topic_embedding = None
    if topic:
        try:
            debate_topic_embedding = embedding_model.encode(topic, convert_to_tensor=False).astype(np.float32)
        except Exception as e:
            logger.error(f"Failed to create embedding for debate topic '{topic}': {e}")
            # Continue without topic embedding if it fails.

    # Build feedback context for each persona based on the topic embedding.
    feedback1_context = build_feedback_prompt_section(user_id, "persona", p1_id, debate_topic_embedding)
    feedback2_context = build_feedback_prompt_section(user_id, "persona", p2_id, debate_topic_embedding)

    # --- Debate Generation ---
    try:
        # Construct the prompt for the debate simulation.
        debate_prompt = (
            f"You are simulating a debate between two distinct AI personas defined by the user.\n"
            f"The topic of the debate is: \"{topic}\"\n\n"
            # Persona 1 Details
            f"Persona 1 (Slot {slot_id_1}): {p1_name}\n"
            f"Instructions for Persona 1:\n--- START P1 INSTRUCTIONS ---\n{p1_prompt}\n--- END P1 INSTRUCTIONS ---\n"
            f"Guidance based on Persona 1's past feedback:\n{feedback1_context}\n\n" # Include feedback
            # Persona 2 Details
            f"Persona 2 (Slot {slot_id_2}): {p2_name}\n"
            f"Instructions for Persona 2:\n--- START P2 INSTRUCTIONS ---\n{p2_prompt}\n--- END P2 INSTRUCTIONS ---\n"
            f"Guidance based on Persona 2's past feedback:\n{feedback2_context}\n\n" # Include feedback
            # Debate Instructions
            f"Generate the debate transcript below. Ensure:\n"
            f"1. Both personas argue according to their instructions, the topic, and any feedback guidance provided.\n"
            f"2. They address each other's points where appropriate.\n"
            f"3. The debate has a clear structure (e.g., opening statement, rebuttals, closing).\n"
            f"4. Clearly label each speaker, like \"{p1_name}: ...\" and \"{p2_name}: ...\".\n"
            f"5. The debate should be comprehensive but stay focused on the topic.\n"
            f"6. Maintain the distinct personality and style of each persona based on their instructions and feedback.\n"
            f"7. Make sure the argument has defintive ending.\n" # Ensure the debate concludes.
            f"Begin the debate now, debate should continue as long as it doesn't reach conclusive end between two parties:"
        )

        logger.debug(f"Generating debate between '{p1_name}' (Slot {slot_id_1}) and '{p2_name}' (Slot {slot_id_2}) on topic '{topic}' for user {user_id}.")
        # Use the dedicated debate model.
        response = await model_debate.generate_content_async(debate_prompt)
        response_text = response.text

        # --- Send Response ---
        # Note: Debate history is not saved to the database by default.
        embed_title = f"🎭 Debate: {p1_name} (Slot {slot_id_1}) vs {p2_name} (Slot {slot_id_2})"
        color_key = "persona_debate"

        # Send response, handling pagination.
        if len(response_text) > config.MAX_EMBED_DESCRIPTION:
            pages = chunk_message(response_text)
            if not pages:
                 logger.error("Chunking resulted in empty pages for debate response.")
                 await send_error_response(interaction, "Failed to process the debate content.")
                 return

            view = PaginatorView(pages, interaction.user, embed_title, color_key)
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, wait=True)
        else:
            embed = discord.Embed(
                title=embed_title,
                description=response_text,
                color=config.EMBED_COLORS.get(color_key, discord.Color.dark_magenta())
            )
            embed.set_footer(text=f"Debate Topic: {topic} | Requested by: {username}", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=embed)

        logger.info(f"Successfully sent debate response for user {user_id}")

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"Debate generation blocked for user {user_id}. Topic: '{topic}'. Reason: {block_err}")
        await send_error_response(interaction, "I cannot generate this debate due to safety restrictions on the topic or persona instructions.")
    except Exception as e:
        logger.exception(f"Error during persona debate generation for user {user_id}, topic '{topic}': {e}")
        await send_error_response(interaction, f"Sorry, an error occurred while generating the debate.")


# --- Utility Command Implementations ---

async def reset_command_impl(interaction: Interaction, conversation_type: str):
    """Implementation for the /reset command."""
    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(ephemeral=True)
    # Determine if resetting all history or a specific type.
    reset_type = None if conversation_type == "all" else conversation_type
    # Call the database function to reset history.
    success, response_message = db.reset_user_history(user_id, reset_type)

    if success:
        await interaction.followup.send(f"✅ {response_message}", ephemeral=True)
    else:
        await send_error_response(interaction, response_message)


async def stats_command_impl(interaction: Interaction):
    """Implementation for the /stats command."""
    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(ephemeral=True)
    # Get user statistics from the database.
    stats = db.get_user_stats(user_id)

    if stats is None:
        logger.warning(f"Could not retrieve stats for user {user_id}.")
        await send_error_response(interaction, "Could not retrieve your statistics. If you're new, try sending a message first!")
        return

    # Create embed to display stats.
    stats_embed = discord.Embed(
        title=f"{username}'s Bot Usage Statistics",
        color=config.EMBED_COLORS.get("user_stats", discord.Color.blue())
    )
    # Format the last reset timestamp.
    try:
         last_reset_str = "Never"
         if stats["last_reset"]:
              last_reset_dt = datetime.datetime.fromisoformat(stats["last_reset"])
              last_reset_str = f"<t:{int(last_reset_dt.timestamp())}:R>" # Relative time format
    except Exception:
         last_reset_str = stats["last_reset"] or "Error parsing date" # Fallback

    # Add fields for each statistic.
    stats_embed.add_field(name="📊 Total Messages Sent (by you)", value=stats.get("total_messages", 0), inline=False)
    # Add counts for each conversation type.
    stats_embed.add_field(name="💬 Nemai (/chat)", value=stats.get("chat_messages", 0), inline=True)
    stats_embed.add_field(name="🕵️ Sherlock (/sherlock)", value=stats.get("sherlock_messages", 0), inline=True)
    stats_embed.add_field(name="👩‍🏫 Teacher (/teacher)", value=stats.get("teacher_messages", 0), inline=True)
    stats_embed.add_field(name="🔬 Scientist (/scientist)", value=stats.get("scientist_messages", 0), inline=True)
    stats_embed.add_field(name="🎭 Personas (/persona chat)", value=stats.get("persona_messages", 0), inline=True)
    stats_embed.add_field(name="📄 Doc Assist (/doc_assist)", value=stats.get("doc_assist_messages", 0), inline=True)
    stats_embed.add_field(name="⭐ Recommend (/recommend)", value=stats.get("recommend_messages", 0), inline=True)
    stats_embed.add_field(name="📖 Story (/story)", value=stats.get("story_messages", 0), inline=True)
    # Add spacer field if needed for alignment (optional)
    # stats_embed.add_field(name="\u200B", value="\u200B", inline=True) # Invisible field

    stats_embed.add_field(name="🔄 Last History Reset", value=last_reset_str, inline=False)
    stats_embed.set_footer(text=f"Requested by: {username}", icon_url=interaction.user.display_avatar.url)
    await interaction.followup.send(embed=stats_embed, ephemeral=True)


async def sentiment_stats_command_impl(interaction: Interaction, conversation_type: str, limit: app_commands.Range[int, 1, config.SENTIMENT_LIMIT_MAX // 2] = config.SENTIMENT_LIMIT_DEFAULT // 2):
    """
    Implementation for the /sentiment_stats command.

    Analyzes the sentiment of the user's recent messages using both a basic
    sentiment pipeline and a more context-aware AI model approach.
    """
    # --- Prerequisite Checks ---
    if not model_chat:
        logger.error("Sentiment analysis requested but chat model is not available.")
        await send_error_response(interaction, "The AI model needed for contextual sentiment analysis is currently unavailable.")
        return
    if not sentiment_pipeline:
        logger.error("Sentiment analysis requested but sentiment pipeline is not available.")
        await send_error_response(interaction, "The basic sentiment analysis tool is currently unavailable.")
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(ephemeral=True)
    # Send an initial message indicating work is in progress.
    initial_message = await interaction.followup.send(f"⏳ Analyzing sentiment for your last {limit} messages (contextual AI, individual labels, and average scores). This may take a moment...", ephemeral=True)

    try:
        # Determine the scope of analysis (all or specific type).
        analysis_type = None if conversation_type == "all" else conversation_type

        # --- Data Retrieval ---
        # Get message pairs (bot context + user response) for contextual analysis.
        message_contexts = db.get_user_message_contexts(user_id, analysis_type, limit)
        # Get individual user messages for basic pipeline analysis.
        user_messages_for_individual = db.get_user_messages(user_id, analysis_type, limit)

        # Check if there's any data to analyze.
        if not message_contexts and not user_messages_for_individual:
            type_str = f"'{conversation_type}'" if analysis_type else "any"
            await interaction.edit_original_response(content=f"I couldn't find any recent messages from you in the {type_str} category to analyze.")
            return

        # --- Contextual Sentiment Analysis (AI Model) ---
        contextual_sentiment_counts = Counter() # Stores counts for each category.
        analyzed_context_count = 0
        contextual_api_errors = 0

        if message_contexts:
            # Define the prompt template for the AI model.
            sentiment_prompt_template = (
                "Analyze the user's sentiment in their message, considering the preceding message from the bot for context. Respond with ONLY ONE of the following sentiment categories, followed by a colon and a brief one-sentence justification:\n\n"
                "Categories:\n"
                "- Appreciative: Expressing thanks or positive acknowledgement.\n"
                "- Positive Engagement: Agreeing happily, confirming positively, showing enthusiasm.\n"
                "- Neutral Agreement/Acknowledgement: Simple confirmation or acknowledgement without strong emotion.\n"
                "- Questioning/Curious: Asking for clarification or more information.\n"
                "- Neutral/Informative: Stating facts, providing information, or message is neutral.\n"
                "- Disagreement/Correction: Disagreeing, correcting, or expressing a negative alternative.\n"
                "- Frustrated/Annoyed: Expressing frustration, annoyance, or dissatisfaction.\n"
                "- Sarcastic/Ironic: Sentiment opposes literal meaning (use sparingly, only if very clear).\n"
                "- Ambiguous/Unclear: Sentiment cannot be determined from the context.\n\n"
                "Bot's Message: \"{bot_context}\"\n"
                "User's Response: \"{user_message}\"\n\n"
                "Sentiment Analysis:"
            )

            logger.info(f"Starting contextual sentiment analysis for user {user_id}, type: {conversation_type}, limit: {len(message_contexts)}. Method: AI Model.")
            # Create asynchronous tasks for each message context.
            tasks = []
            for context in message_contexts:
                prompt = sentiment_prompt_template.format(
                    bot_context=context['bot_context'][:1000], # Truncate context if needed
                    user_message=context['user_message'][:1000] # Truncate message if needed
                )
                tasks.append(model_chat.generate_content_async(prompt))

            # Run tasks concurrently and gather results.
            contextual_results = await asyncio.gather(*tasks, return_exceptions=True)

            # Define the expected categories from the prompt.
            valid_contextual_categories = {
                "Appreciative", "Positive Engagement", "Neutral Agreement/Acknowledgement",
                "Questioning/Curious", "Neutral/Informative", "Disagreement/Correction",
                "Frustrated/Annoyed", "Sarcastic/Ironic", "Ambiguous/Unclear"
            }

            # Process the results from the AI.
            for result in contextual_results:
                if isinstance(result, Exception):
                    logger.error(f"Error during contextual sentiment API call: {result}")
                    contextual_api_errors += 1
                    continue
                try:
                    response_text = result.text.strip()
                    sentiment_category = "Ambiguous/Unclear" # Default if parsing fails.
                    # Extract the category before the first colon.
                    if ':' in response_text:
                        extracted_category = response_text.split(':', 1)[0].strip()
                        # Validate if the extracted category is one we expect.
                        if extracted_category in valid_contextual_categories:
                            sentiment_category = extracted_category
                        else:
                            logger.warning(f"LLM returned unexpected contextual sentiment category: {extracted_category}")
                    else:
                         logger.warning(f"LLM contextual sentiment response format invalid: {response_text}")
                    # Increment the count for the determined category.
                    contextual_sentiment_counts[sentiment_category] += 1
                    analyzed_context_count += 1
                except Exception as e:
                    # Handle errors processing individual responses.
                    logger.error(f"Error processing contextual sentiment response '{getattr(result, 'text', 'N/A')}': {e}")
                    contextual_sentiment_counts["Ambiguous/Unclear"] += 1 # Count as unclear on error.
                    contextual_api_errors +=1
            logger.info(f"Contextual sentiment analysis complete for user {user_id}. Analyzed: {analyzed_context_count}. Errors: {contextual_api_errors}. Counts: {dict(contextual_sentiment_counts)}")
        else:
            logger.info(f"No message contexts found for user {user_id}, type {conversation_type}. Skipping contextual analysis.")

        # --- Individual Message Sentiment Analysis (Transformers Pipeline) ---
        individual_sentiment_counts = Counter() # Positive, Negative, Neutral, Unknown
        individual_sentiment_scores = {'Positive': 0.0, 'Negative': 0.0, 'Neutral': 0.0} # Sum of scores
        individual_sentiment_successful_analyses = {'Positive': 0, 'Negative': 0, 'Neutral': 0} # Count for averaging
        analyzed_individual_count = 0
        individual_analysis_errors = 0

        if user_messages_for_individual:
            logger.info(f"Starting individual sentiment analysis for user {user_id}, type: {conversation_type}, limit: {len(user_messages_for_individual)}. Method: Transformers Pipeline.")
            for user_message in user_messages_for_individual:
                # Use the basic sentiment analysis function.
                sentiment_result = get_sentiment_analysis(user_message)
                analyzed_individual_count += 1
                if sentiment_result:
                    label = sentiment_result.get('label')
                    score = sentiment_result.get('score', 0.0)
                    # Categorize and aggregate scores.
                    # Note: The specific labels ('POSITIVE', 'NEGATIVE') depend on the model used.
                    if label == 'POSITIVE':
                        individual_sentiment_counts['Positive'] += 1
                        individual_sentiment_scores['Positive'] += score
                        individual_sentiment_successful_analyses['Positive'] += 1
                    elif label == 'NEGATIVE':
                        individual_sentiment_counts['Negative'] += 1
                        individual_sentiment_scores['Negative'] += score
                        individual_sentiment_successful_analyses['Negative'] += 1
                    else: # Assume any other label is Neutral for simplicity.
                        individual_sentiment_counts['Neutral'] += 1
                        individual_sentiment_scores['Neutral'] += score
                        individual_sentiment_successful_analyses['Neutral'] += 1
                else:
                    # Count messages where basic analysis failed.
                    individual_sentiment_counts['Unknown'] += 1
                    individual_analysis_errors += 1
            logger.info(f"Individual sentiment analysis complete for user {user_id}. Analyzed: {analyzed_individual_count} attempts. Successes: {sum(individual_sentiment_successful_analyses.values())}. Errors: {individual_analysis_errors}. Counts: {dict(individual_sentiment_counts)}")
        else:
            logger.info(f"No individual messages found for user {user_id}, type {conversation_type}. Skipping individual analysis.")

        # --- Result Aggregation and Display ---
        # Check if any analysis was successful.
        if analyzed_context_count == 0 and analyzed_individual_count == 0:
            final_content = "❌ Failed to analyze sentiment. No messages found or all analysis methods encountered issues."
            if contextual_api_errors > 0: final_content += f" ({contextual_api_errors} contextual errors)"
            if individual_analysis_errors > 0: final_content += f" ({individual_analysis_errors} individual errors)"
            await interaction.edit_original_response(content=final_content)
            return

        # Create the result embed.
        type_display = conversation_type.capitalize() if analysis_type else "All"
        embed = discord.Embed(
            title=f"{username}'s Sentiment Analysis",
            description=(
                f"Analysis based on up to {limit} of your recent messages.\n"
                f"*(Conversation Type: {type_display})*"
            ),
            color=config.EMBED_COLORS.get("sentiment_stats")
        )

        # Add Contextual Analysis Results field.
        if analyzed_context_count > 0:
            total_categorized_contextual = analyzed_context_count
            # Define display order and emojis for contextual categories.
            contextual_category_details = [
                ("Appreciative", "😊", contextual_sentiment_counts["Appreciative"]),
                ("Positive Engagement", "👍", contextual_sentiment_counts["Positive Engagement"]),
                ("Neutral Agreement/Acknowledgement", "👌", contextual_sentiment_counts["Neutral Agreement/Acknowledgement"]),
                ("Questioning/Curious", "🤔", contextual_sentiment_counts["Questioning/Curious"]),
                ("Neutral/Informative", "😐", contextual_sentiment_counts["Neutral/Informative"]),
                ("Disagreement/Correction", "👎", contextual_sentiment_counts["Disagreement/Correction"]),
                ("Frustrated/Annoyed", "😠", contextual_sentiment_counts["Frustrated/Annoyed"]),
                ("Sarcastic/Ironic", "😏", contextual_sentiment_counts["Sarcastic/Ironic"]),
                ("Ambiguous/Unclear", "❓", contextual_sentiment_counts["Ambiguous/Unclear"]),
            ]
            contextual_distribution_text = ""
            # Format the distribution string.
            for name, emoji, count in contextual_category_details:
                if count > 0:
                     percentage = round((count / total_categorized_contextual) * 100, 1)
                     contextual_distribution_text += f"{emoji} {name}: {count} ({percentage}%)\n"
            if not contextual_distribution_text: contextual_distribution_text = "No messages were categorized by the AI."
            embed.add_field(
                name=f"📊 AI Contextual Sentiment (Analyzed: {analyzed_context_count})",
                value=contextual_distribution_text.strip(),
                inline=False
            )
            # Add error count if any occurred.
            if contextual_api_errors > 0: embed.add_field(name="⚠️ AI Analysis Issues", value=f"Could not analyze {contextual_api_errors} messages due to API/processing issues.", inline=False)
        elif message_contexts: # If contexts existed but analysis failed
             embed.add_field(name="📊 AI Contextual Sentiment", value=f"Contextual analysis failed or no messages were categorized (Errors: {contextual_api_errors}).", inline=False)

        # Add Individual Analysis Results fields.
        total_individual_analyzed_success = sum(individual_sentiment_successful_analyses.values())
        if analyzed_individual_count > 0:
            ind_pos_count = individual_sentiment_counts['Positive']
            ind_neg_count = individual_sentiment_counts['Negative']
            ind_neu_count = individual_sentiment_counts['Neutral']
            ind_unk_count = individual_sentiment_counts['Unknown']

            # Calculate percentages.
            ind_pos_perc = round((ind_pos_count / analyzed_individual_count) * 100, 1) if analyzed_individual_count > 0 else 0
            ind_neg_perc = round((ind_neg_count / analyzed_individual_count) * 100, 1) if analyzed_individual_count > 0 else 0
            ind_neu_perc = round((ind_neu_count / analyzed_individual_count) * 100, 1) if analyzed_individual_count > 0 else 0
            ind_unk_perc = round((ind_unk_count / analyzed_individual_count) * 100, 1) if analyzed_individual_count > 0 else 0

            # Format distribution text.
            individual_distribution_text = f"😊 Positive: {ind_pos_count} ({ind_pos_perc}%)\n"
            individual_distribution_text += f"☹️ Negative: {ind_neg_count} ({ind_neg_perc}%)\n"
            individual_distribution_text += f"😐 Neutral: {ind_neu_count} ({ind_neu_perc}%)"
            if ind_unk_count > 0: individual_distribution_text += f"\n❓ Unknown: {ind_unk_count} ({ind_unk_perc}%)"

            # Determine the dominant sentiment label based on counts.
            dominant_sentiment = "Neutral"
            dominant_emoji = "😐"
            counts_for_dominant = {'Positive': ind_pos_count, 'Negative': ind_neg_count, 'Neutral': ind_neu_count}
            max_count = 0
            if counts_for_dominant:
                max_count = max(counts_for_dominant.values())

            if max_count > 0:
                dominant_candidates = [k for k, v in counts_for_dominant.items() if v == max_count]
                if len(dominant_candidates) == 1: # Single dominant sentiment
                    dominant_sentiment = dominant_candidates[0]
                    if dominant_sentiment == "Positive": dominant_emoji = "😊"
                    elif dominant_sentiment == "Negative": dominant_emoji = "☹️"
                elif len(dominant_candidates) > 1: # Multiple sentiments tied
                    dominant_sentiment = "Mixed"
                    dominant_emoji = "🤔"
            elif ind_unk_count > 0: # If only unknown results
                 dominant_sentiment = "Unknown"
                 dominant_emoji = "❓"

            # Add fields for individual label distribution and dominant label.
            embed.add_field(
                name=f"📈 Individual Message Labels (Analyzed: {analyzed_individual_count})",
                value=individual_distribution_text,
                inline=True # Display side-by-side with dominant label
            )
            embed.add_field(
                name=f"{dominant_emoji} Dominant Individual Label",
                value=f"{dominant_sentiment}",
                inline=True
            )

            # Calculate and add average scores if any messages were successfully analyzed.
            avg_pos_score = (individual_sentiment_scores['Positive'] / individual_sentiment_successful_analyses['Positive']) if individual_sentiment_successful_analyses['Positive'] > 0 else 0.0
            avg_neg_score = (individual_sentiment_scores['Negative'] / individual_sentiment_successful_analyses['Negative']) if individual_sentiment_successful_analyses['Negative'] > 0 else 0.0
            avg_neu_score = (individual_sentiment_scores['Neutral'] / individual_sentiment_successful_analyses['Neutral']) if individual_sentiment_successful_analyses['Neutral'] > 0 else 0.0

            if total_individual_analyzed_success > 0:
                average_score_text = (f"😊 Positive Avg. Score: {avg_pos_score:.3f} (from {individual_sentiment_successful_analyses['Positive']} msgs)\n"
                                      f"☹️ Negative Avg. Score: {avg_neg_score:.3f} (from {individual_sentiment_successful_analyses['Negative']} msgs)\n"
                                      f"😐 Neutral Avg. Score: {avg_neu_score:.3f} (from {individual_sentiment_successful_analyses['Neutral']} msgs)")
                embed.add_field(
                    name=f"⚖️ Average Individual Scores (Analyzed: {total_individual_analyzed_success})",
                    value=average_score_text,
                    inline=False
                )

            # Add error count if any occurred.
            if individual_analysis_errors > 0:
                embed.add_field(name="⚠️ Individual Analysis Issues", value=f"Could not analyze {individual_analysis_errors} individual messages.", inline=False)

        elif user_messages_for_individual: # If messages existed but analysis failed
             embed.add_field(name="📈 Individual Message Analysis", value=f"Individual analysis failed or no messages were analyzed (Errors: {individual_analysis_errors}).", inline=False)

        # Finalize and send the embed.
        embed.set_footer(text="Sentiment analysis complete.") # Updated footer
        await interaction.edit_original_response(content=None, embed=embed) # Edit the initial "⏳ Analyzing..." message.

    except Exception as e:
        # Catch-all for unexpected errors during the command execution.
        logger.exception(f"Error during sentiment analysis command for user {user_id}: {e}")
        try:
             # Try to edit the original message with an error.
             await interaction.edit_original_response(content="An unexpected error occurred while analyzing your sentiment.")
        except discord.NotFound: pass # Ignore if original message is gone.
        except discord.HTTPException:
             # If editing fails, send a new followup message.
             await interaction.followup.send("An unexpected error occurred while analyzing your sentiment.", ephemeral=True)


async def help_command_impl(interaction: Interaction):
    """Implementation for the /help command."""
    await interaction.response.defer(ephemeral=True) # Help is usually ephemeral.

    # Create the main help embed.
    help_embed = discord.Embed(
        title="🤖 Nemai Bot Help",
        description="Hello! I'm Nemai, your AI assistant. Here's what I can do:",
        color=config.EMBED_COLORS.get("default")
    )
    # Set the bot's avatar as the thumbnail.
    if client and client.user:
        help_embed.set_thumbnail(url=client.user.display_avatar.url)

    # Add fields for different command categories.
    help_embed.add_field(
        name="💬 Core Chat Commands",
        value="`/chat [message]`: General chat with Nemai.\n"
              "`/sherlock [message]`: Consult with Sherlock Holmes.\n"
              "`/teacher [message]`: Get simple explanations.\n"
              "`/scientist [message]`: Discuss with a scientist persona.",
        inline=False
    )
    help_embed.add_field(
        name="📰🌐 File, Web & Utility Commands",
        value="`/search [query]`: Search the web (via DuckDuckGo) and get an AI summary.\n"
              "`/news [topic]`: Get AI-summarized news for a topic (or general news if no topic).\n"
              "`/factcheck [statement]`: Verify a statement using web search and AI analysis.\n"
              "`/analyze_file [file] [prompt]`: Analyze an uploaded image or text file. Ask questions using the optional prompt.\n"
              "`/search_history [query] [limit]`: Search your past messages for relevant info (private).\n"
              "`/explain_like_im [topic] [audience]`: Explain a topic based on web search, tailored to a specific audience.\n"
              "`/summarize [text_or_url]`: Summarize provided text or the content of a web page (via search context).\n"
              "`/recipe [dish_name]`: Find and format a recipe for a specific dish using web search.",
        inline=False
    )

    # Document Assistance Commands (Grouped)
    help_embed.add_field(
        name=f"📄 Document Assistance (`/doc_assist ...`)",
        value="`/doc_assist analyze [type] [file]`: Get initial AI analysis of a resume, cover letter, essay, or report.\n"
              "`/doc_assist critique [file] [focus]`: Get specific feedback (clarity, grammar, impact, conciseness) on an analyzed file.\n"
              "`/doc_assist rewrite_section [file] [goal] [section_text]`: Rewrite a section of an analyzed file (professional, concise, engaging).",
        inline=False
    )

    # Recommendation Commands (Grouped)
    help_embed.add_field(
        name=f"⭐ Recommendation Engine (`/recommend ...`)",
        value="`/recommend get [type] [genre] [based_on]`: Get recommendations (movie, book, music, game).\n"
              "`/recommend set_preference [type] [likes] [dislikes]`: Set your preferences for future recommendations.",
        inline=False
    )

    # Story Generator Commands (Grouped)
    help_embed.add_field(
        name=f"📖 Story Generator (`/story ...`)",
        value="`/story start [mode] [genre] [setting]`: Begin a new story (collaborative or choose_your_own).\n"
              "`/story continue [your_turn]`: Continue a collaborative story.\n"
              "`/story end`: Finish your current active story.\n"
              "`/story status`: Check your current story's progress.",
        inline=False
    )

    # Special Request Feature (Classification)
    help_embed.add_field(
        name="✨ Special Request Features (Text Commands)",
        value="Within `/chat`, `/teacher`, or `/scientist` you can ask for:\n"
              "• **Definitions:** e.g., 'define photosynthesis'\n"
              "• **Comparisons:** e.g., 'compare python and javascript'\n"
              "• **Citations:** e.g., 'find a citation for quantum entanglement'\n"
              "• **Pros & Cons:** e.g., 'what are the pros and cons of electric cars'",
        inline=False
    )

    # Custom Persona Commands (Grouped)
    help_embed.add_field(
        name=f"🎭 Custom Persona Commands (`/persona ...`)",
        value=f"*(Max {config.MAX_PERSONAS} personas per user, managed in Slots 1-{config.MAX_PERSONAS})*\n"
              "`... create [name] [description]`: Create a new persona (assigns to next free slot).\n"
              "`... activate [slot_id]`: Make a persona active for chatting.\n"
              "`... chat [message]`: Chat with your *active* persona.\n"
              "`... debate [slot1] [slot2] [topic]`: Make two personas debate a topic. *(History not saved)*\n"
              "`... list`: Show all your personas and their slots.\n"
              "`... info [slot_id]`: View details and prompt of a persona.\n"
              "`... delete [slot_id]`: Remove a persona from its slot.",
        inline=False
    )

    # Account/History Commands
    help_embed.add_field(
        name="🔧 Account & History Commands",
        value="`/stats`: View your message counts.\n"
              "`/sentiment_stats [type] [limit]`: Analyze the sentiment of your recent messages.\n"
              "`/reset [type]`: Clear conversation history.\n"
              "`/export [type] [limit]`: Export history to a file.\n"
              "`/help`: Show this help message.",
        inline=False
    )

    # Add info about the feedback system if enabled.
    if config.ENABLE_FEEDBACK_SYSTEM:
         help_embed.add_field(
             name="👍👎 Feedback System",
             value="You'll see thumbs up/down buttons on my responses. Clicking them helps me learn what you like and dislike, especially for similar future messages!",
             inline=False
         )

    help_embed.set_footer(text=f"Nemai Bot | Use commands to interact!")
    await interaction.followup.send(embed=help_embed, ephemeral=True)


async def export_command_impl(interaction: Interaction, conversation_type: str, limit: int):
    """Implementation for the /export command."""
    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(ephemeral=True) # Export is private.

    # Determine scope (all or specific type).
    export_type = None if conversation_type == "all" else conversation_type

    try:
        # Retrieve history tuples (role, content, timestamp) from the database.
        history_tuples = db.get_conversation_history_for_export(user_id, export_type, limit)

        if not history_tuples:
            type_str = f"'{conversation_type}'" if export_type else "any"
            await interaction.followup.send(f"❌ No history found in the {type_str} category to export.", ephemeral=True)
            return

        # --- Format Export Content ---
        export_content = f"Conversation History Export\n"
        export_content += f"User: {username} ({user_id})\n"
        export_content += f"Type: {conversation_type.capitalize() if export_type else 'All'}\n"
        export_content += f"Limit: {limit} messages\n"
        export_content += f"Exported At: {datetime.datetime.now(datetime.timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}\n"
        export_content += "-----------------------------------------\n\n"

        # Add each message to the export string.
        for role, content, timestamp_str in history_tuples:
            # Format timestamp.
            try:
                dt_obj = datetime.datetime.fromisoformat(timestamp_str)
                ts_formatted = dt_obj.strftime("%Y-%m-%d %H:%M:%S")
            except ValueError:
                ts_formatted = "Unknown Time" # Fallback
            export_content += f"[{ts_formatted}] {role}: {content}\n\n" # Add extra newline for readability.

        export_content += "-----------------------------------------\nEnd of Export"

        # --- Create and Send File ---
        # Encode the content to bytes and create an in-memory file.
        file_content = io.BytesIO(export_content.encode('utf-8'))
        # Create a unique filename.
        filename = f"nemai_export_{user_id}_{conversation_type}_{datetime.datetime.now().strftime('%Y%m%d%H%M')}.txt"
        export_file = discord.File(fp=file_content, filename=filename)

        # Send the file ephemerally.
        await interaction.followup.send(f"✅ Here is your conversation history export ({len(history_tuples)} messages):", file=export_file, ephemeral=True)
        logger.info(f"Successfully generated and sent history export for user {user_id}, type {conversation_type}, limit {limit}.")

    except Exception as e:
        logger.exception(f"Error during history export for user {user_id}: {e}")
        await send_error_response(interaction, "An error occurred while exporting your history.")


# --- Doc Assist Command Implementations ---

async def doc_assist_analyze_impl(
    interaction: Interaction,
    doc_type: Literal['resume', 'cover_letter', 'essay', 'report'],
    file: Attachment
):
    """Implementation for the /doc_assist analyze command."""
    # Reuse the generic file analysis logic with a specific prompt.
    await analyze_file_command_impl(interaction, file, prompt=f"Analyze this {doc_type.replace('_', ' ')}.")


async def doc_assist_critique_impl(
    interaction: Interaction,
    doc_type: Literal['resume', 'cover_letter', 'essay', 'report'],
    file: Attachment,
    focus: Literal['clarity', 'grammar', 'impact', 'conciseness']
):
    """Implementation for the /doc_assist critique command."""
    if not model_chat:
        await send_error_response(interaction, "AI model unavailable for critique.")
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    # --- File Validation ---
    file_type = file.content_type
    if file_type not in config.ALLOWED_TEXT_TYPES: # Critique only works on text.
        allowed_types_str = ", ".join(sorted(list(set([t.split('/')[-1] for t in config.ALLOWED_TEXT_TYPES]))))
        await send_error_response(interaction, f"Unsupported file type ('{file_type}') for critique. Please use: {allowed_types_str}.")
        return
    if file.size > config.MAX_FILE_SIZE_BYTES:
         await send_error_response(interaction, f"File is too large ({file.size / (1024*1024):.2f} MB). Maximum allowed size is {config.MAX_FILE_SIZE_BYTES / (1024*1024):.2f} MB.")
         return

    await interaction.response.defer(thinking=True)

    try:
        # --- Text Extraction ---
        file_bytes = await file.read()
        file_content_str = None
        try:
            # Extract text using the same logic as analyze_file.
            if file_type == "application/pdf":
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                file_content_str = "".join([page.extract_text() for page in pdf_reader.pages])
            elif file_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                doc = docx.Document(io.BytesIO(file_bytes))
                file_content_str = "\n".join([para.text for para in doc.paragraphs])
            else:
                 try: file_content_str = file_bytes.decode('utf-8')
                 except UnicodeDecodeError: file_content_str = file_bytes.decode('latin-1')
        except Exception as parse_err:
             logger.error(f"Failed to parse file '{file.filename}' for critique: {parse_err}", exc_info=True)
             await send_error_response(interaction, f"Could not extract text from '{file.filename}' for critique.")
             return

        if not file_content_str or not file_content_str.strip():
             await send_error_response(interaction, f"Could not find any readable text content in '{file.filename}'.")
             return

        # Truncate if necessary.
        if len(file_content_str) > config.MAX_FILE_CONTENT_CHARS:
            file_content_str = file_content_str[:config.MAX_FILE_CONTENT_CHARS]
            logger.warning(f"File content for '{file.filename}' truncated for critique.")

        # --- AI Critique ---
        # Construct the prompt for targeted critique.
        critique_prompt = (
            f"You are an expert writing assistant. Provide a constructive critique of the following {doc_type.replace('_', ' ')} "
            f"focusing specifically on **{focus}**. Analyze the provided text and offer specific, actionable suggestions for improvement "
            f"related to {focus}. Point out both strengths and weaknesses related to the focus area.\n\n"
            f"--- DOCUMENT CONTENT ({file.filename}) START ---\n"
            f"{file_content_str}\n"
            f"--- DOCUMENT CONTENT ({file.filename}) END ---\n\n"
            f"Critique (Focus: {focus}):"
        )

        response = await model_chat.generate_content_async(critique_prompt)
        response_text = response.text

        # --- Save and Send ---
        embed_title = f"📄 Critique ({focus.capitalize()}): {file.filename}"
        color_key = "doc_assist"

        # Save interaction to DB (optional, but good practice).
        db.save_message(user_id, "doc_assist", "User", f"Critique request: {doc_type}, {file.filename}, focus: {focus}")
        bot_msg_id, _ = db.save_message(user_id, "doc_assist", "Nemai", response_text) # Save bot response

        # Send response, handling pagination.
        if len(response_text) > config.MAX_EMBED_DESCRIPTION:
            pages = chunk_message(response_text)
            view = PaginatorView(pages, interaction.user, embed_title, color_key)
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, wait=True)
        else:
            embed = discord.Embed(title=embed_title, description=response_text, color=config.EMBED_COLORS.get(color_key))
            embed.set_footer(text=f"Requested by: {username}", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=embed)

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"Doc critique blocked for user {user_id}. Reason: {block_err}")
        await send_error_response(interaction, "Critique failed due to safety restrictions.")
    except Exception as e:
        logger.exception(f"Error during doc critique for user {user_id}: {e}")
        await send_error_response(interaction, "An error occurred while generating the critique.")


async def doc_assist_rewrite_section_impl(
    interaction: Interaction,
    doc_type: Literal['resume', 'cover_letter', 'essay', 'report'],
    goal: Literal['make more professional', 'make more concise', 'make more engaging'],
    section_text: str
):
    """Implementation for the /doc_assist rewrite_section command."""
    if not model_chat:
        await send_error_response(interaction, "AI model unavailable for rewriting.")
        return
    if not section_text or not section_text.strip():
        await send_error_response(interaction, "Please provide the text section to rewrite.")
        return
    # Limit input length to avoid overly long requests/responses.
    if len(section_text) > config.DOC_ASSIST_MAX_SECTION_LEN:
         await send_error_response(interaction, f"Section text is too long (max {config.DOC_ASSIST_MAX_SECTION_LEN} characters). Please shorten it.")
         return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(thinking=True)

    try:
        # --- AI Rewrite ---
        # Construct the prompt for rewriting the section with a specific goal.
        rewrite_prompt = (
            f"You are an expert writing assistant. Rewrite the following text section, which is part of a {doc_type.replace('_', ' ')}, "
            f"to achieve the goal: **{goal}**. Maintain the original meaning but adjust the tone, style, and wording as needed "
            f"to meet the specified goal. Provide only the rewritten section.\n\n"
            f"--- ORIGINAL SECTION START ---\n"
            f"{section_text}\n"
            f"--- ORIGINAL SECTION END ---\n\n"
            f"Rewritten Section (Goal: {goal}):"
        )

        response = await model_chat.generate_content_async(rewrite_prompt)
        response_text = response.text.strip() # Remove leading/trailing whitespace

        # --- Save and Send ---
        embed_title = f"📄 Rewritten Section ({goal.capitalize()})"
        color_key = "doc_assist"

        # Save interaction to DB.
        db.save_message(user_id, "doc_assist", "User", f"Rewrite request: {doc_type}, goal: {goal}, original: {section_text[:100]}...")
        bot_msg_id, _ = db.save_message(user_id, "doc_assist", "Nemai", response_text)

        # Format the response to show both original and rewritten text.
        full_description = f"**Original:**\n```\n{section_text}\n```\n**Rewritten ({goal}):**\n{response_text}"

        # Send response, handling pagination.
        if len(full_description) > config.MAX_EMBED_DESCRIPTION:
            # If combined text is too long, paginate the full description.
            pages = chunk_message(full_description)
            view = PaginatorView(pages, interaction.user, embed_title, color_key)
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, wait=True)
        else:
            embed = discord.Embed(title=embed_title, description=full_description, color=config.EMBED_COLORS.get(color_key))
            embed.set_footer(text=f"Requested by: {username}", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=embed)

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"Doc rewrite blocked for user {user_id}. Reason: {block_err}")
        await send_error_response(interaction, "Rewrite failed due to safety restrictions.")
    except Exception as e:
        logger.exception(f"Error during doc rewrite for user {user_id}: {e}")
        await send_error_response(interaction, "An error occurred while rewriting the section.")


# --- Recommendation Command Implementations ---

async def recommend_set_preference_impl(
    interaction: Interaction,
    pref_type: Literal['movie', 'book', 'music', 'game'],
    likes: Optional[str] = None,
    dislikes: Optional[str] = None
):
    """Implementation for the /recommend set_preference command."""
    # Ensure at least one preference is provided.
    if not likes and not dislikes:
        await send_error_response(interaction, "Please provide at least likes or dislikes.", ephemeral=True)
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(ephemeral=True) # Preferences are private.
    # Save preferences to the database.
    success, message = db.set_user_preference(user_id, pref_type, likes, dislikes)

    if success:
        await interaction.followup.send(f"✅ {message}", ephemeral=True)
    else:
        await send_error_response(interaction, f"Failed to set preferences: {message}", ephemeral=True)


async def recommend_get_impl(
    interaction: Interaction,
    rec_type: Literal['movie', 'book', 'music', 'game'],
    genre: Optional[str] = None,
    based_on: Optional[str] = None
):
    """Implementation for the /recommend get command."""
    if not model_chat:
        await send_error_response(interaction, "AI model unavailable for recommendations.")
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(thinking=True)

    # --- Gather Context ---
    # Retrieve stored user preferences for this recommendation type.
    preferences = db.get_user_preferences(user_id, rec_type)
    likes, dislikes = preferences if preferences else (None, None)

    # Format preferences for the prompt.
    pref_text = ""
    if likes: pref_text += f"\n- User Likes: {likes}"
    if dislikes: pref_text += f"\n- User Dislikes: {dislikes}"
    if pref_text: pref_text = f"\nConsider the user's stored preferences:{pref_text}"

    # Format the user's specific request details.
    request_text = f"Recommend {config.RECOMMENDATION_LIMIT} {rec_type}s."
    if genre: request_text += f" In the {genre} genre."
    if based_on: request_text += f" Similar to or based on: {based_on}."

    # --- Optional Web Search for Recent Items ---
    search_context = ""
    # Construct a search query based on user input.
    search_query = f"new popular {genre if genre else ''} {rec_type}s {f'like {based_on}' if based_on else ''}".strip()
    try:
        # Perform a quick web search for potentially relevant recent items.
        loop = asyncio.get_running_loop()
        # Limit results to keep context concise.
        results = await loop.run_in_executor(None, lambda: list(DDGS().text(search_query, max_results=3)))
        if results:
            search_context += "\n\nRecent relevant items found online (for context, prioritize user request/prefs):\n"
            for result in results:
                # Add snippets from search results.
                search_context += f"- {result.get('title', 'N/A')}: {result.get('body', 'N/A')[:100]}...\n"
    except Exception as search_err:
        # Don't fail the command if search fails, just log a warning.
        logger.warning(f"Web search for recommendation context failed: {search_err}")

    # --- AI Recommendation ---
    # Construct the prompt for the AI.
    recommend_prompt = (
        f"You are a helpful recommendation assistant. {request_text}\n"
        f"{pref_text}\n{search_context}\n" # Include preferences and search context.
        f"Provide a list of {config.RECOMMENDATION_LIMIT} specific {rec_type} recommendations. "
        f"For each recommendation, give a brief (1-2 sentence) justification explaining why the user might like it, "
        f"connecting it to their request or preferences if possible. Format as a numbered list."
    )

    try:
        response = await model_chat.generate_content_async(recommend_prompt)
        response_text = response.text

        # --- Save and Send ---
        embed_title = f"⭐ {rec_type.capitalize()} Recommendations"
        color_key = "recommend"

        # Save interaction to DB.
        db.save_message(user_id, "recommend", "User", f"Recommend request: {rec_type}, genre: {genre}, based_on: {based_on}")
        bot_msg_id, _ = db.save_message(user_id, "recommend", "Nemai", response_text)

        # Send response, handling pagination.
        if len(response_text) > config.MAX_EMBED_DESCRIPTION:
            pages = chunk_message(response_text)
            view = PaginatorView(pages, interaction.user, embed_title, color_key)
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, wait=True)
        else:
            embed = discord.Embed(title=embed_title, description=response_text, color=config.EMBED_COLORS.get(color_key))
            embed.set_footer(text=f"Requested by: {username}", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=embed)

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"Recommendation blocked for user {user_id}. Reason: {block_err}")
        await send_error_response(interaction, "Recommendation failed due to safety restrictions.")
    except Exception as e:
        logger.exception(f"Error during recommendation generation for user {user_id}: {e}")
        await send_error_response(interaction, "An error occurred while generating recommendations.")


# --- Story Generator Command Implementations & View ---

async def story_process_choice_impl(interaction: Interaction, session_id: int, choice_text: str):
    """
    Handles the logic when a user clicks a choice button in a 'choose_your_own' story.

    Args:
        interaction: The button click interaction.
        session_id: The ID of the story session.
        choice_text: The text of the choice the user selected.
    """
    user_id = str(interaction.user.id)
    username = interaction.user.display_name # Not strictly needed here, but good practice

    # --- Session Validation ---
    # Get the active story session details from the database.
    active_session = db.get_story_session_details(session_id, user_id)

    if not active_session:
        await send_error_response(interaction, "Could not find your active story session for this choice.", ephemeral=True)
        return

    # Unpack session details.
    _, genre, setting, mode, story_state_json, _ = active_session

    # Ensure this function is only called for the correct story mode.
    if mode != 'choose_your_own':
        await send_error_response(interaction, "This action is only valid for 'choose your own' stories.", ephemeral=True)
        return

    # --- Update Story State ---
    try:
        # Load the story history from the JSON string.
        story_history = json.loads(story_state_json)
    except json.JSONDecodeError:
        logger.error(f"Could not decode story state JSON for session {session_id} during choice processing.")
        await send_error_response(interaction, "Error loading story history.", ephemeral=True)
        db.end_story_session(session_id) # End the corrupted session.
        return

    # Add the user's choice as a turn in the history.
    user_choice_turn = f"User chose to: {choice_text}"
    # Update the session in the database first.
    db.update_story_session(session_id, {"role": "User", "content": user_choice_turn})
    # Also save it as a regular message (optional, for general history).
    db.save_message(user_id, "story", "User", user_choice_turn)
    # Append to the local history list for the prompt.
    story_history.append({"role": "User", "content": user_choice_turn})

    # --- Generate Next Story Part ---
    try:
        # Create context from the most recent turns.
        context = "\n".join([f"{turn['role']}: {turn['content']}" for turn in story_history[-config.STORY_CONTEXT_LIMIT:]])
        # Construct the prompt for the AI to continue the story and provide new choices.
        continue_prompt = (
            f"Continue the story based on the user's choice and the preceding context.\n"
            f"Story Context:\n{context}\n\n"
            f"Mode: {mode}. Genre: {genre}. Setting: {setting}.\n"
            f"Write the next part of the story from the AI's perspective, following the user's choice. Keep it consistent and engaging.\n"
            f"Then, present 2-3 new distinct choices (A, B, C...). Make sure the choices are clearly separated from the narrative.\n" # Instruct AI to provide choices.
            f"Next part:"
        )

        response = await model_chat.generate_content_async(continue_prompt)
        next_part = response.text.strip()

        # Update the database with the AI's turn.
        db.update_story_session(session_id, {"role": "AI", "content": next_part})
        bot_msg_id, _ = db.save_message(user_id, "story", "Nemai", next_part) # Save to general history too.

        # --- Prepare and Send Response with Choices ---
        embed_title = f"📖 Story Time ({mode.replace('_', ' ').title()})"
        color_key = "story"
        embed = discord.Embed(title=embed_title, description=next_part, color=config.EMBED_COLORS.get(color_key))
        embed.set_footer(text=f"Story Session ID: {session_id}")

        # --- Extract Choices from AI Response ---
        # Attempt to parse choices (A., B., C. or 1., 2., 3. etc.) from the end of the AI's response.
        choices = []
        lines = next_part.split('\n')
        # Look at the last few lines for potential choices.
        possible_choice_lines = [line.strip() for line in lines[-5:] if line.strip()]
        choice_started = False
        for line in possible_choice_lines:
            # Check for common choice prefixes or if we're already parsing choices and the line is short.
            if line.startswith(('A.', 'B.', 'C.', '1.', '2.', '3.', '- ','* ')) or (choice_started and len(line) < 100):
                # Extract text after the prefix.
                choice_text_part = line.split('.', 1)[-1].strip().lstrip('-* ')
                if choice_text_part:
                    choices.append(choice_text_part)
                    choice_started = True # Assume subsequent short lines are also choices.
            else:
                 choice_started = False # Reset if a non-choice line is encountered.

        choices = [c for c in choices if c] # Remove any empty choices extracted.
        # Provide fallback choices if parsing fails or yields too few options.
        if len(choices) < 2:
            logger.warning(f"Could not extract sufficient choices from story part for session {session_id}. Using fallbacks.")
            choices = ["Continue the adventure.", "Look around."]
        else:
             choices = choices[:3] # Limit to a maximum of 3 choices.

        # Create the view with the extracted choices.
        view = StoryChoiceView(session_id, user_id, choices)

        # Disable buttons on the *previous* message (the one the user just clicked).
        if interaction.message:
            try:
                 await interaction.message.edit(view=None) # Remove the old view.
            except (discord.NotFound, discord.Forbidden, discord.HTTPException) as e:
                 logger.warning(f"Could not disable buttons on previous story message {interaction.message.id}: {e}")

        # Send the new story part with the new choice buttons.
        # Use followup since the interaction was deferred in the button callback.
        followup_message = await interaction.followup.send(embed=embed, view=view)
        view.message = followup_message # Store message for timeout handling.

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"Story continuation blocked for session {session_id} after choice. Reason: {block_err}")
        await send_error_response(interaction, "Cannot continue story due to safety restrictions.", ephemeral=True)
        # Consider ending the session here? db.end_story_session(session_id)
    except Exception as e:
        logger.exception(f"Error processing story choice for session {session_id}: {e}")
        await send_error_response(interaction, "An error occurred while continuing the story after your choice.", ephemeral=True)
        # Consider ending the session here? db.end_story_session(session_id)


class StoryChoiceView(ui.View):
    """
    A Discord UI View containing buttons for story choices in 'choose_your_own' mode.

    Attributes:
        session_id: The ID of the story session these choices belong to.
        user_id: The Discord ID of the user who can interact with these choices.
        selected_choice: Stores the text of the selected choice (not currently used after selection).
        message: The discord.Message object this view is attached to.
    """
    def __init__(self, session_id: int, user_id: str, choices: List[str]):
        """
        Initializes the StoryChoiceView with buttons for each provided choice.

        Args:
            session_id: The story session ID.
            user_id: The ID of the user allowed to interact.
            choices: A list of strings representing the choices.
        """
        # Set a timeout for the choices.
        super().__init__(timeout=config.STORY_TIMEOUT_MINUTES * 60)
        self.session_id = session_id
        self.user_id = user_id
        self.selected_choice: Optional[str] = None
        self.message: Optional[discord.Message] = None

        # Dynamically create a button for each choice.
        for i, choice_text in enumerate(choices):
            # Create a unique custom_id for each button.
            button = ui.Button(
                label=choice_text[:80], # Truncate long choice text for button label.
                style=ButtonStyle.secondary,
                custom_id=f"story_choice_{session_id}_{i}" # Include session_id and index.
            )
            # Assign the callback using a factory function to capture the correct choice_text.
            button.callback = self.create_callback(choice_text)
            self.add_item(button)

    def create_callback(self, choice_text: str):
        """
        Factory function to create a button callback that captures the specific choice text.

        This is necessary because button callbacks are defined in the class scope,
        but we need each button instance to know which choice it represents.

        Args:
            choice_text: The text of the choice this callback is for.

        Returns:
            An asynchronous callback function for the button.
        """
        async def button_callback(interaction: Interaction):
            # Check if the interacting user is the allowed user.
            if str(interaction.user.id) != self.user_id:
                await interaction.response.send_message("This story choice isn't for you.", ephemeral=True)
                return

            # Defer the response immediately as processing takes time.
            # thinking=True shows the "Bot is thinking..." indicator.
            await interaction.response.defer(thinking=True)

            # Disable all buttons in this view after a choice is made.
            for item in self.children:
                if isinstance(item, ui.Button):
                    item.disabled = True
            # Edit the message the button was attached to, removing the view (or showing disabled buttons).
            if interaction.message: # Check if message still exists
                try:
                    await interaction.message.edit(view=self) # Show disabled buttons
                except (discord.NotFound, discord.HTTPException):
                    pass # Ignore if message is gone

            # Call the main logic function to process the chosen action.
            await story_process_choice_impl(interaction, self.session_id, choice_text)

            # Stop listening for interactions on this specific view instance.
            self.stop()
        return button_callback

    async def on_timeout(self):
        """Disables buttons when the view times out."""
        self.selected_choice = None # Indicate no choice was made.
        if self.message:
            try:
                # Check if buttons are already disabled.
                already_disabled = all(item.disabled for item in self.children if isinstance(item, ui.Button))
                if not already_disabled:
                    for item in self.children:
                         if isinstance(item, ui.Button):
                             item.disabled = True
                    await self.message.edit(view=self) # Edit message to show disabled buttons.
            except (discord.NotFound, discord.Forbidden, discord.HTTPException):
                # Ignore errors if message is gone or permissions are missing.
                pass
        self.stop() # Stop the view listener.


async def story_start_impl(interaction: Interaction, mode: Literal['collaborative', 'choose_your_own'], genre: Optional[str] = None, setting: Optional[str] = None):
     """Implementation for the /story start command."""
     user_id = str(interaction.user.id)
     username = interaction.user.display_name
     db.add_user(user_id, username)

     # --- End Previous Session (If Any) ---
     # Ensure a user only has one active story session at a time.
     active_session = db.get_active_story_session(user_id)
     if active_session:
          # End the previous session before starting a new one.
          db.end_story_session(active_session[0])
          logger.info(f"Ended previous active story session {active_session[0]} for user {user_id} before starting new one.")

     await interaction.response.defer(thinking=True)

     # --- Create New Session ---
     # Create the new story session entry in the database.
     session_id, error = db.create_story_session(user_id, genre, setting, mode)
     if error:
          await send_error_response(interaction, f"Failed to start story: {error}")
          return

     # --- Generate First Part ---
     try:
          # Construct the prompt for the AI to start the story.
          start_prompt = (
               f"Start a new story. Mode: {mode}. Genre: {genre if genre else 'freestyle'}. Setting: {setting if setting else 'open'}.\n"
               f"Write the first paragraph of the story. Keep it engaging and open-ended.\n"
               # Add instruction to provide choices only for 'choose_your_own' mode.
               f"{'Then, present 2-3 distinct choices for the user to guide the next step (Label choices clearly, e.g., A., B., C. or 1., 2., 3. on separate lines after the narrative). Make sure choices are clearly separated from narrative text.' if mode == 'choose_your_own' else ''}"
               f"First paragraph:"
          )
          response = await model_chat.generate_content_async(start_prompt)
          first_part = response.text.strip()

          # --- Save and Send ---
          # Save the initial user request and the AI's first part to the database.
          db.save_message(user_id, "story", "User", f"Start Story: {mode}, Genre: {genre}, Setting: {setting}")
          db.update_story_session(session_id, {"role": "AI", "content": first_part}) # Add AI turn to session state.
          bot_msg_id, _ = db.save_message(user_id, "story", "Nemai", first_part) # Save to general history.

          # Prepare the embed for the first part.
          embed_title = f"📖 Story Time ({mode.replace('_', ' ').title()})"
          color_key = "story"
          embed = discord.Embed(title=embed_title, description=first_part, color=config.EMBED_COLORS.get(color_key))
          embed.set_footer(text=f"Story started by: {username} | Session ID: {session_id}")

          # --- Add Choice Buttons (If Applicable) ---
          view = None
          if mode == 'choose_your_own':
               # Extract choices from the AI's response using the same logic as story_process_choice_impl.
               choices = []
               lines = first_part.split('\n')
               possible_choice_lines = [line.strip() for line in lines[-5:] if line.strip()]
               choice_started = False
               for line in possible_choice_lines:
                   if line.startswith(('A.', 'B.', 'C.', '1.', '2.', '3.', '- ','* ')) or (choice_started and len(line) < 100):
                       choice_text_part = line.split('.', 1)[-1].strip().lstrip('-* ')
                       if choice_text_part:
                           choices.append(choice_text_part)
                           choice_started = True
                   else:
                        choice_started = False

               choices = [c for c in choices if c]
               if len(choices) < 2:
                   logger.warning(f"Could not extract sufficient choices from story start for session {session_id}. Using fallbacks.")
                   choices = ["Continue the adventure.", "Look around."]
               else:
                    choices = choices[:3]

               # Create the view with the extracted choices.
               view = StoryChoiceView(session_id, user_id, choices)

          # Send the first story part, including the choice view if created.
          followup_message = await interaction.followup.send(embed=embed, view=view if view else discord.utils.MISSING)
          if view: view.message = followup_message # Store message in view for timeout handling.

     except genai.types.BlockedPromptException as block_err:
          logger.warning(f"Story start blocked for user {user_id}. Reason: {block_err}")
          if session_id: db.end_story_session(session_id) # Clean up the created session.
          await send_error_response(interaction, "Failed to start story due to safety restrictions.")
     except TypeError as te: # Catch potential type errors during processing
          logger.exception(f"TypeError during story start for user {user_id}: {te}")
          if session_id: db.end_story_session(session_id)
          await send_error_response(interaction, f"A type error occurred while starting the story: {te}")
     except Exception as e:
          logger.exception(f"Error starting story for user {user_id}: {e}")
          if session_id: db.end_story_session(session_id) # Clean up on any error.
          await send_error_response(interaction, "An error occurred while starting the story.")


async def story_continue_impl(interaction: Interaction, your_turn: Optional[str] = None):
     """Implementation for the /story continue command (collaborative mode)."""
     user_id = str(interaction.user.id)
     username = interaction.user.display_name
     db.add_user(user_id, username)

     await interaction.response.defer(thinking=True)

     # --- Session Validation ---
     active_session = db.get_active_story_session(user_id)
     if not active_session:
          await send_error_response(interaction, "You don't have an active story session. Use `/story start`.")
          return

     session_id, genre, setting, mode, story_state_json, _ = active_session

     # Ensure this command is used only for collaborative mode.
     if mode != 'collaborative':
         await send_error_response(interaction, "Use `/story start` or choice buttons for 'choose your own' mode. `/story continue` is for collaborative mode only.")
         return

     # Ensure the user provided their part of the story.
     if not your_turn or not your_turn.strip():
          await send_error_response(interaction, "Please provide your part of the story for collaborative mode.")
          return

     # --- Update Story State ---
     try:
          story_history = json.loads(story_state_json)
     except json.JSONDecodeError:
          logger.error(f"Could not decode story state JSON for session {session_id}")
          await send_error_response(interaction, "Error loading story history. Please start a new story.")
          db.end_story_session(session_id) # End corrupted session.
          return

     # Add the user's turn to the history.
     user_input = your_turn
     db.update_story_session(session_id, {"role": "User", "content": user_input})
     db.save_message(user_id, "story", "User", user_input) # Save to general history.
     story_history.append({"role": "User", "content": user_input})

     # --- Generate AI's Next Turn ---
     try:
          # Create context and prompt for the AI.
          context = "\n".join([f"{turn['role']}: {turn['content']}" for turn in story_history[-config.STORY_CONTEXT_LIMIT:]])
          continue_prompt = (
               f"Continue the story based on the last user input and the preceding context.\n"
               f"Story Context:\n{context}\n\n"
               f"Mode: {mode}. Genre: {genre}. Setting: {setting}.\n"
               f"Write the next part of the story from the AI's perspective. Keep it consistent and engaging.\n" # No choices needed for collaborative.
               f"Next part:"
          )

          response = await model_chat.generate_content_async(continue_prompt)
          next_part = response.text.strip()

          # --- Save and Send ---
          db.update_story_session(session_id, {"role": "AI", "content": next_part})
          bot_msg_id, _ = db.save_message(user_id, "story", "Nemai", next_part)

          embed_title = f"📖 Story Time ({mode.replace('_', ' ').title()})"
          color_key = "story"
          embed = discord.Embed(title=embed_title, description=next_part, color=config.EMBED_COLORS.get(color_key))
          embed.set_footer(text=f"Story Session ID: {session_id}")

          # Send the AI's turn.
          await interaction.followup.send(embed=embed)

     except genai.types.BlockedPromptException as block_err:
          logger.warning(f"Story continuation blocked for session {session_id}. Reason: {block_err}")
          await send_error_response(interaction, "Cannot continue story due to safety restrictions.")
     except Exception as e:
          logger.exception(f"Error continuing story session {session_id}: {e}")
          await send_error_response(interaction, "An error occurred while continuing the story.")


async def story_status_impl(interaction: Interaction):
    """Implementation for the /story status command."""
    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)
    await interaction.response.defer(ephemeral=True) # Status is private.

    # Get the user's active story session.
    active_session = db.get_active_story_session(user_id)
    if not active_session:
        await interaction.followup.send("You don't have an active story session.", ephemeral=True)
        return

    # Unpack details.
    session_id, genre, setting, mode, story_state_json, last_updated_str = active_session
    # Get a snippet of the last turn.
    try:
        story_history = json.loads(story_state_json)
        last_turn = story_history[-1]['content'] if story_history else "No turns yet."
    except:
        last_turn = "[Error loading history]" # Fallback

    # Format last updated time.
    try: last_updated_dt = datetime.datetime.fromisoformat(last_updated_str)
    except: last_updated_dt = None
    last_updated_display = f"<t:{int(last_updated_dt.timestamp())}:R>" if last_updated_dt else "Unknown"

    # Create and send the status embed.
    embed = discord.Embed(title=f"📖 Active Story Status (Session {session_id})", color=config.EMBED_COLORS.get("story"))
    embed.add_field(name="Mode", value=mode.replace('_', ' ').title(), inline=True)
    embed.add_field(name="Genre", value=genre if genre else "N/A", inline=True)
    embed.add_field(name="Setting", value=setting if setting else "N/A", inline=True)
    embed.add_field(name="Last Update", value=last_updated_display, inline=False)
    # Show snippet in a code block.
    embed.add_field(name="Last Turn Snippet", value=f"```{last_turn[:1000]}...```", inline=False)
    embed.set_footer(text="Use /story continue (for collaborative) or buttons (for choose_your_own), or /story end.")

    await interaction.followup.send(embed=embed, ephemeral=True)


async def story_end_impl(interaction: Interaction):
    """Implementation for the /story end command."""
    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)
    await interaction.response.defer(ephemeral=True) # Confirmation is private.

    # Find the active session.
    active_session = db.get_active_story_session(user_id)
    if not active_session:
        await interaction.followup.send("You don't have an active story session to end.", ephemeral=True)
        return

    # End the session in the database.
    session_id = active_session[0]
    if db.end_story_session(session_id):
        await interaction.followup.send(f"✅ Story session {session_id} ended.", ephemeral=True)
    else:
        # This might happen if the session timed out between check and end.
        await send_error_response(interaction, f"Could not end story session {session_id}. It might already be inactive.", ephemeral=True)


async def recipe_command_impl(interaction: Interaction, dish_name: str):
    """
    Implementation for the /recipe command.

    Finds a recipe using web search and formats it using the AI model.
    """
    if not model_chat:
        logger.error("Cannot handle recipe: Chat model not available.")
        await send_error_response(interaction, "The AI model needed for recipe generation is currently unavailable.")
        return
    if not dish_name.strip():
        await send_error_response(interaction, "Please provide the name of the dish you want a recipe for.")
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(thinking=True)
    logger.info(f"User {user_id} requested recipe for: '{dish_name}'")

    # --- Web Search ---
    # Construct a search query likely to yield recipe results.
    search_query = f"{dish_name} recipe simple ingredients instructions"
    search_results_text = ""
    try:
        loop = asyncio.get_running_loop()
        results = await loop.run_in_executor(
            None,
            lambda: list(DDGS().text(search_query, max_results=config.SEARCH_RESULT_LIMIT))
        )

        if not results:
            logger.warning(f"No search results found for recipe query: '{search_query}'")
            await interaction.followup.send(f"❌ No web search results found for a '{dish_name}' recipe. Try a different name or phrasing?")
            return

        # Format results for the AI prompt.
        search_results_text += f"Search Results for Recipe Query: \"{search_query}\"\n\n"
        for i, result in enumerate(results):
            title = result.get('title', 'No Title')
            body = result.get('body', 'No Snippet')
            href = result.get('href', 'No URL')
            search_results_text += f"Result {i+1}: {title}\nSnippet: {body}\nURL: <{href}>\n\n"

        logger.debug(f"Fetched {len(results)} results for recipe query: '{search_query}'")

    except Exception as e:
        logger.exception(f"Error performing DuckDuckGo search for recipe query '{search_query}': {e}")
        await send_error_response(interaction, "Sorry, I encountered an error while searching the web for recipes.")
        return

    # --- AI Recipe Formatting ---
    try:
        # Construct the prompt to extract and format the recipe.
        recipe_prompt = (
            f"Based *only* on the following web search results provided for a \"{dish_name}\" recipe, "
            f"compile a clear and easy-to-follow recipe. Synthesize the information from the snippets.\n\n"
            f"Instructions:\n"
            f"1. Create two distinct sections: 'Ingredients' and 'Instructions'.\n"
            f"2. List all necessary ingredients clearly under the 'Ingredients' heading. Use bullet points.\n"
            f"3. Provide step-by-step instructions under the 'Instructions' heading. Use numbered points.\n"
            f"4. Extract and combine information from the search results. Do not add external knowledge or ingredients/steps not mentioned.\n"
            f"5. If the results are contradictory or insufficient to form a complete recipe, try to synthesize a coherent recipe from the most promising result(s), or state that a reliable recipe could not be formed.\n" # Updated instruction 5
            f"6. Aim for a standard, common version of the recipe based on the search results.\n"
            f"7. Be concise and practical.\n\n"
            f"--- Search Results ---\n"
            f"{search_results_text}\n"
            f"--- End Search Results ---\n\n"
            f"Recipe for {dish_name}:" # Start the recipe directly
        )

        logger.debug(f"Generating recipe for: '{dish_name}'")
        response = await model_chat.generate_content_async(recipe_prompt)
        recipe_text = response.text.strip()

        # Basic validation: check if key sections are present.
        if not recipe_text or "ingredients" not in recipe_text.lower() or "instructions" not in recipe_text.lower():
             logger.warning(f"AI generated an incomplete or invalid recipe for: '{dish_name}'. Response: {recipe_text[:200]}...")
             # Provide a more informative error message to the user.
             recipe_text = f"I couldn't construct a reliable recipe for '{dish_name}' based on the search results. The information might have been insufficient or conflicting."

        # --- Send Response ---
        embed_title = f"🍳 Recipe: {dish_name.title()}" # Title case the dish name.
        color_key = "recipe"

        # Send response, handling pagination.
        if len(recipe_text) > config.MAX_EMBED_DESCRIPTION:
            pages = chunk_message(recipe_text)
            if not pages:
                 logger.error("Chunking resulted in empty pages for recipe.")
                 await send_error_response(interaction, "Failed to process the recipe content.")
                 return

            view = PaginatorView(pages, interaction.user, embed_title, color_key)
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, wait=True)
        else:
            embed = discord.Embed(
                title=embed_title,
                description=recipe_text,
                color=config.EMBED_COLORS.get(color_key, discord.Color.blue())
            )
            embed.set_footer(text=f"Recipe requested by: {username} | Info synthesized from web search", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=embed)

        logger.info(f"Successfully sent recipe for '{dish_name}' to user {user_id}")

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"Recipe generation blocked for '{dish_name}'. Reason: {block_err}")
        await send_error_response(interaction, "I cannot generate a recipe for this dish due to safety restrictions.")
    except Exception as e:
        logger.exception(f"Error generating recipe for '{dish_name}': {e}")
        await send_error_response(interaction, "Sorry, I encountered an error while generating the recipe.")


async def summarize_command_impl(interaction: Interaction, text_or_url: str):
    """
    Implementation for the /summarize command.

    Summarizes either a block of text or the content of a URL (fetched via web search context).
    """
    if not model_chat:
        logger.error("Cannot handle summarize: Chat model not available.")
        await send_error_response(interaction, "The AI model needed for summarization is currently unavailable.")
        return
    if not text_or_url.strip():
        await send_error_response(interaction, "Please provide text or a URL to summarize.")
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username)

    await interaction.response.defer(thinking=True)

    content_to_summarize = ""
    source_type = "text"
    source_display = "Provided Text"
    # Basic regex to check if the input looks like a URL.
    url_pattern = re.compile(r'https?://\S+')
    is_url = url_pattern.match(text_or_url)

    # --- Get Content to Summarize ---
    if is_url:
        # If input is a URL, fetch context using web search.
        url = text_or_url
        source_type = "url"
        source_display = f"URL: <{url}>" # Format for Discord link
        logger.info(f"User {user_id} requested summary for URL: '{url}'")

        search_results_text = ""
        try:
            loop = asyncio.get_running_loop()
            # First, try searching specifically for the site and URL.
            results = await loop.run_in_executor(
                None,
                lambda: list(DDGS().text(f'site:{url} {url}', max_results=config.SEARCH_RESULT_LIMIT // 2)) # Limit results
            )
            # If site-specific search fails, try searching for the URL directly.
            if not results:
                results = await loop.run_in_executor(
                    None,
                    lambda: list(DDGS().text(url, max_results=config.SEARCH_RESULT_LIMIT // 2)) # Limit results
                )

            if not results:
                logger.warning(f"No search results found for URL context: '{url}'")
                await interaction.followup.send(f"❌ Could not find web context for the URL '{url}'. Unable to summarize.")
                return

            # Format search results as context.
            search_results_text += f"Web Context for URL: \"{url}\"\n\n"
            for i, result in enumerate(results):
                title = result.get('title', 'No Title')
                body = result.get('body', 'No Snippet')
                href = result.get('href', 'No URL')
                search_results_text += f"Context {i+1}: {title}\nSnippet: {body}\nURL: <{href}>\n\n"

            content_to_summarize = search_results_text
            logger.debug(f"Fetched {len(results)} context results for URL: '{url}'")

        except Exception as e:
            logger.exception(f"Error performing DuckDuckGo search for URL context '{url}': {e}")
            await send_error_response(interaction, "Sorry, I encountered an error while fetching context for the URL.")
            return

    else:
        # If input is not a URL, treat it as raw text.
        source_type = "text"
        source_display = "Provided Text"
        content_to_summarize = text_or_url
        logger.info(f"User {user_id} requested summary for provided text (length: {len(content_to_summarize)}).")
        # Truncate very long input text.
        max_text_len = 10000 # Define a reasonable max length for direct text summary.
        if len(content_to_summarize) > max_text_len:
            content_to_summarize = content_to_summarize[:max_text_len]
            logger.warning(f"Input text for summary truncated to {max_text_len} characters.")

    # --- AI Summarization ---
    try:
        summarization_prompt = ""
        # Create different prompts based on whether it's a URL or text.
        if source_type == "url":
            summarization_prompt = (
                f"Based *only* on the following web search context provided for the URL \"{url}\", "
                f"write a concise and informative summary of what the URL is likely about. "
                f"Synthesize the key information found in the snippets. "
                f"Do not add external knowledge. If the context seems irrelevant or contradictory, mention that briefly.\n\n"
                f"--- Web Context ---\n"
                f"{content_to_summarize}\n"
                f"--- End Web Context ---\n\n"
                f"Summary:"
            )
        else: # source_type == "text"
            summarization_prompt = (
                f"Please provide a concise summary of the following text:\n\n"
                f"--- Text Start ---\n"
                f"{content_to_summarize}\n"
                f"--- Text End ---\n\n"
                f"Summary:"
            )

        logger.debug(f"Generating summary for {source_type}: '{text_or_url[:100]}...'")
        response = await model_chat.generate_content_async(summarization_prompt)
        summary_text = response.text.strip()

        if not summary_text:
             logger.warning(f"AI generated an empty summary for {source_type}: '{text_or_url[:100]}...'")
             summary_text = "The AI couldn't generate a summary based on the provided content."

        # --- Send Response ---
        embed_title = f"📝 Summary"
        color_key = "summary"

        # Combine source info and summary for the description.
        full_description = f"**Source:** {source_display}\n\n**Summary:**\n{summary_text}"

        # Handle pagination if the combined description is too long.
        if len(full_description) > config.MAX_EMBED_DESCRIPTION:
            # Paginate only the summary part.
            summary_pages = chunk_message(summary_text)
            if not summary_pages:
                 logger.error("Chunking resulted in empty pages for summary.")
                 await send_error_response(interaction, "Failed to process the summary content.")
                 return

            # Prepend the source info to the first page.
            first_page_content = f"**Source:** {source_display}\n\n**Summary:**\n{summary_pages[0]}"
            # Check if even the first page with source info is too long.
            if len(first_page_content) > config.MAX_EMBED_DESCRIPTION:
                # If so, just use the first chunk of the summary and adjust title.
                first_page_content = summary_pages[0]
                embed_title = f"📝 Summary for {source_display}" # Add source to title instead.

            # Combine the potentially modified first page with the rest.
            final_pages = [first_page_content] + summary_pages[1:]

            view = PaginatorView(final_pages, interaction.user, embed_title, color_key)
            view.message = await interaction.followup.send(embed=view.get_page_embed(), view=view, wait=True)
        else:
            # Send as single embed if it fits.
            embed = discord.Embed(
                title=embed_title,
                description=full_description,
                color=config.EMBED_COLORS.get(color_key, discord.Color.blue())
            )
            embed.set_footer(text=f"Summary requested by: {username}", icon_url=interaction.user.display_avatar.url)
            await interaction.followup.send(embed=embed)

        logger.info(f"Successfully sent summary to user {user_id} for {source_type}.")

    except genai.types.BlockedPromptException as block_err:
        logger.warning(f"Summary generation blocked for {source_type} '{text_or_url[:100]}...'. Reason: {block_err}")
        await send_error_response(interaction, "I cannot generate a summary for this content due to safety restrictions.")
    except Exception as e:
        logger.exception(f"Error summarizing {source_type} '{text_or_url[:100]}...': {e}")
        await send_error_response(interaction, "Sorry, I encountered an error while generating the summary.")


async def imagine_command_impl(interaction: Interaction, prompt: str):
    """
    Implementation for the /imagine command.

    Generates an image using a Hugging Face Inference API endpoint (e.g., Stable Diffusion).
    Requires HUGGINGFACE_TOKEN to be configured.
    """
    # --- Prerequisite Checks ---
    if not config.HUGGINGFACE_TOKEN:
        logger.error("Imagine command invoked but HUGGINGFACE_TOKEN is not configured.")
        await send_error_response(interaction, "Image generation is currently disabled by the bot administrator.", ephemeral=True)
        return
    if not prompt or not prompt.strip():
        await send_error_response(interaction, "Please provide a prompt describing the image you want to create.", ephemeral=True)
        return

    user_id = str(interaction.user.id)
    username = interaction.user.display_name
    db.add_user(user_id, username) # Ensure user exists

    await interaction.response.defer(thinking=True)
    logger.info(f"User {user_id} initiated image generation with prompt: '{prompt[:100]}...'")

    # --- API Request ---
    headers = {"Authorization": f"Bearer {config.HUGGINGFACE_TOKEN}", "Content-Type": "application/json"}
    payload = {"inputs": prompt}
    # Set a timeout for the API request.
    api_timeout = aiohttp.ClientTimeout(total=config.IMAGE_GEN_TIMEOUT_SECONDS)

    try:
        # Use aiohttp for asynchronous HTTP requests.
        async with aiohttp.ClientSession(timeout=api_timeout) as session:
            async with session.post(config.IMAGE_GEN_MODEL_URL, headers=headers, json=payload) as response:
                # --- Handle API Response ---
                if response.status == 200:
                    # Check if the response content type is an image.
                    if 'image' in response.content_type:
                        image_bytes = await response.read()
                        # Create a discord.File object from the image bytes.
                        image_file = discord.File(io.BytesIO(image_bytes), filename="generated_image.png") # Assume png, can be adjusted

                        # Create embed to display the image and prompt.
                        embed = discord.Embed(
                            title=f"🎨 Image Generation Complete",
                            description=f"Prompt: `{prompt}`",
                            color=config.EMBED_COLORS.get("imagine", discord.Color.pink())
                        )
                        # Attach the image file to the embed.
                        embed.set_image(url=f"attachment://{image_file.filename}")
                        embed.set_footer(text=f"Requested by: {username}", icon_url=interaction.user.display_avatar.url)

                        # Send the embed and the image file.
                        await interaction.followup.send(embed=embed, file=image_file)
                        logger.info(f"Successfully sent generated image for prompt '{prompt[:50]}...' to user {user_id}")

                    else:
                        # Handle cases where API returns 200 OK but not an image.
                        response_text = await response.text()
                        logger.error(f"Hugging Face API returned status 200 but content type was not image ({response.content_type}). Response: {response_text[:200]}")
                        await send_error_response(interaction, "Image generation failed: The API returned an unexpected response format.")
                        return

                elif response.status == 503: # Service Unavailable (often means model is loading)
                    try:
                        response_json = await response.json()
                        estimated_time = response_json.get('estimated_time', 0)
                        logger.warning(f"Hugging Face API returned 503 (Service Unavailable/Model Loading). Estimated time: {estimated_time}s. Prompt: '{prompt[:50]}...'")
                        # Inform user about the estimated wait time.
                        await send_error_response(interaction, f"The image generation model is currently loading. Please try again in {int(estimated_time) + 5} seconds.")
                    except (json.JSONDecodeError, aiohttp.ContentTypeError):
                         logger.warning(f"Hugging Face API returned 503 but response was not valid JSON. Prompt: '{prompt[:50]}...'")
                         await send_error_response(interaction, "The image generation model is currently loading or unavailable. Please try again shortly.")
                    return
                elif response.status == 429: # Rate Limited
                     logger.warning(f"Hugging Face API returned 429 (Rate Limited). User: {user_id}")
                     await send_error_response(interaction, "Image generation failed: Too many requests. Please wait a moment and try again.")
                     return
                else:
                    # Handle other API errors.
                    response_text = await response.text()
                    logger.error(f"Hugging Face API request failed with status {response.status}. Response: {response_text[:500]}. Prompt: '{prompt[:50]}...'")
                    await send_error_response(interaction, f"Image generation failed. The API returned status code {response.status}. Please check the prompt or try again later.")
                    return

    except asyncio.TimeoutError:
        # Handle request timeout.
        logger.error(f"Hugging Face API request timed out after {config.IMAGE_GEN_TIMEOUT_SECONDS}s for prompt: '{prompt[:50]}...'")
        await send_error_response(interaction, f"Image generation timed out after {config.IMAGE_GEN_TIMEOUT_SECONDS} seconds. The request might be too complex or the service is busy.")
    except aiohttp.ClientConnectorError as e:
         # Handle network connection errors.
         logger.error(f"Network connection error during image generation: {e}")
         await send_error_response(interaction, "Could not connect to the image generation service. Please check network or try again later.")
    except Exception as e:
        # Catch any other unexpected errors.
        logger.exception(f"An unexpected error occurred during image generation for user {user_id}, prompt '{prompt[:50]}...': {e}")
        await send_error_response(interaction, "An unexpected error occurred while generating the image.")


# --- Command Setup and Registration ---

def setup_commands(
    _db: Database,
    _client: discord.Client,
    _tree: app_commands.CommandTree,
    _persona_group: app_commands.Group,
    _doc_assist_group: app_commands.Group,
    _recommend_group: app_commands.Group,
    _story_group: app_commands.Group,
    _model_chat: Optional[genai.GenerativeModel],
    _model_persona: Optional[genai.GenerativeModel],
    _model_vision: Optional[genai.GenerativeModel],
    _model_debate: Optional[genai.GenerativeModel],
    _sentiment_pipeline: Optional[pipeline],
    _embedding_model: Optional[Any], # Using Any as type hint for flexibility
    _logger: logging.Logger
):
    """
    Initializes global variables and registers all slash commands with Discord.

    This function is called once when the bot starts up (in `on_ready`).
    It assigns the necessary objects (database, client, models, logger) to the
    global variables in this module and then defines and registers each
    application command (slash command) and command group.

    Args:
        _db: The initialized Database instance.
        _client: The initialized discord.Client instance.
        _tree: The app_commands.CommandTree instance associated with the client.
        _persona_group: The app_commands.Group for persona commands.
        _doc_assist_group: The app_commands.Group for document assistance commands.
        _recommend_group: The app_commands.Group for recommendation commands.
        _story_group: The app_commands.Group for story generator commands.
        _model_chat: The initialized Gemini model for general chat.
        _model_persona: The initialized Gemini model for personas/debate.
        _model_vision: The initialized Gemini model for vision tasks.
        _model_debate: The initialized Gemini model specifically configured for debates.
        _sentiment_pipeline: The initialized Hugging Face sentiment analysis pipeline.
        _embedding_model: The initialized sentence embedding model.
        _logger: The main logger instance.
    """
    global db, client, tree, persona_group, doc_assist_group, recommend_group, story_group
    global model_chat, model_persona, model_vision, model_debate, sentiment_pipeline, embedding_model, logger

    # Assign passed objects to global variables.
    db = _db
    client = _client
    tree = _tree
    persona_group = _persona_group
    doc_assist_group = _doc_assist_group
    recommend_group = _recommend_group
    story_group = _story_group
    model_chat = _model_chat
    model_persona = _model_persona
    model_vision = _model_vision
    model_debate = _model_debate
    sentiment_pipeline = _sentiment_pipeline
    embedding_model = _embedding_model
    logger = _logger

    # --- Command Registration ---
    # The following decorators register the functions as slash commands.
    # The `_reg` suffix distinguishes the registration function from the implementation function.

    # Core Chat Commands
    @tree.command(name="chat", description="Chat with Nemai (general AI assistant)")
    @app_commands.describe(message="Your message or question for Nemai.")
    async def chat_command_reg(interaction: discord.Interaction, message: str):
        """Registration wrapper for the /chat command."""
        await chat_command_impl(interaction, message)

    @tree.command(name="sherlock", description="Chat with the detective Sherlock Holmes")
    @app_commands.describe(message="Present your case or question to Sherlock.")
    async def sherlock_command_reg(interaction: discord.Interaction, message: str):
        """Registration wrapper for the /sherlock command."""
        await sherlock_command_impl(interaction, message)

    @tree.command(name="teacher", description="Ask for explanations in a simple, teacher-like style")
    @app_commands.describe(message="What concept do you want the teacher to explain?")
    async def teacher_command_reg(interaction: discord.Interaction, message: str):
        """Registration wrapper for the /teacher command."""
        await teacher_command_impl(interaction, message)

    @tree.command(name="scientist", description="Discuss topics with a knowledgeable scientist persona")
    @app_commands.describe(message="Your question or topic for the scientist.")
    async def scientist_command_reg(interaction: discord.Interaction, message: str):
        """Registration wrapper for the /scientist command."""
        await scientist_command_impl(interaction, message)

    # Web/Utility Commands
    @tree.command(name="search", description="Search the web (via DuckDuckGo) and get an AI-summarized answer.")
    @app_commands.describe(query="What do you want to search for?")
    async def search_command_reg(interaction: discord.Interaction, query: str):
        """Registration wrapper for the /search command."""
        await search_command_impl(interaction, query)

    @tree.command(name="news", description="Get summarized recent news headlines for a topic or general news.")
    @app_commands.describe(topic="Optional: Specific topic for news (e.g., 'technology', 'finance'). Leave blank for top headlines.")
    async def news_command_reg(interaction: discord.Interaction, topic: str | None = None):
        """Registration wrapper for the /news command."""
        await news_command_impl(interaction, topic)

    @tree.command(name="analyze_file", description="Analyze an image or text-based file")
    @app_commands.describe(file="The image or text file to analyze.", prompt="Optional text prompt or question about the file.")
    async def analyze_file_command_reg(interaction: discord.Interaction, file: discord.Attachment, prompt: str | None = None):
        """Registration wrapper for the /analyze_file command."""
        await analyze_file_command_impl(interaction, file, prompt)

    @tree.command(name="factcheck", description="Fact-check a statement using web search and AI analysis.")
    @app_commands.describe(statement="The statement you want to verify.")
    async def factcheck_command_reg(interaction: discord.Interaction, statement: str):
        """Registration wrapper for the /factcheck command."""
        await factcheck_command_impl(interaction, statement)

    @tree.command(name="search_history", description="Search your past conversation history with the bot (private).")
    @app_commands.describe(query="What keywords or topic are you looking for in your history?", limit=f"Max number of results (default {config.HISTORY_SEARCH_RESULT_LIMIT})")
    async def history_search_command_reg(interaction: discord.Interaction, query: str, limit: app_commands.Range[int, 1, 50] = config.HISTORY_SEARCH_RESULT_LIMIT):
        """Registration wrapper for the /search_history command."""
        await history_search_command_impl(interaction, query, limit)

    @tree.command(name="explain_like_im", description="Explain a topic based on web search, tailored to a specific audience.")
    @app_commands.describe(topic="The subject to explain.", audience="Describe the audience (e.g., a 5-year-old, a physics professor).")
    async def explain_like_im_reg(interaction: discord.Interaction, topic: str, audience: str):
        """Registration wrapper for the /explain_like_im command."""
        await explain_like_im_impl(interaction, topic, audience)

    @tree.command(name="summarize", description="Summarize provided text or the content of a URL (via web context).")
    @app_commands.describe(text_or_url="The text block or a valid web URL to summarize.")
    async def summarize_command_reg(interaction: discord.Interaction, text_or_url: str):
        """Registration wrapper for the /summarize command."""
        await summarize_command_impl(interaction, text_or_url)

    @tree.command(name="recipe", description="Find a recipe for a dish using web search and AI formatting.")
    @app_commands.describe(dish_name="The name of the dish (e.g., 'chocolate chip cookies', 'chicken curry').")
    async def recipe_command_reg(interaction: discord.Interaction, dish_name: str):
        """Registration wrapper for the /recipe command."""
        await recipe_command_impl(interaction, dish_name)

    # Persona Management Commands (Grouped)
    @persona_group.command(name="create", description=f"Create a new custom persona in the next available slot (1-{config.MAX_PERSONAS})")
    @app_commands.describe(name=f"A unique name for your persona ({config.PERSONA_NAME_MIN_LEN}-{config.PERSONA_NAME_MAX_LEN} chars).", description=f"Detailed instructions for the persona's behavior ({config.PERSONA_DESC_MIN_LEN}-{config.PERSONA_DESC_MAX_LEN} chars).")
    async def persona_create_reg(interaction: discord.Interaction, name: str, description: str):
        """Registration wrapper for the /persona create command."""
        await persona_create_impl(interaction, name, description)

    @persona_group.command(name="delete", description="Delete one of your custom personas by its slot number")
    @app_commands.describe(slot_id=f"The slot number (1-{config.MAX_PERSONAS}) of the persona to delete.")
    async def persona_delete_reg(interaction: discord.Interaction, slot_id: app_commands.Range[int, 1, config.MAX_PERSONAS]):
        """Registration wrapper for the /persona delete command."""
        await persona_delete_impl(interaction, slot_id)

    @persona_group.command(name="list", description="List all of your created personas and their slots (1-5)")
    async def persona_list_reg(interaction: discord.Interaction):
        """Registration wrapper for the /persona list command."""
        await persona_list_impl(interaction)

    @persona_group.command(name="info", description="Show details about one of your personas by its slot number")
    @app_commands.describe(slot_id=f"The slot number (1-{config.MAX_PERSONAS}) of the persona to view.")
    async def persona_info_reg(interaction: discord.Interaction, slot_id: app_commands.Range[int, 1, config.MAX_PERSONAS]):
        """Registration wrapper for the /persona info command."""
        await persona_info_impl(interaction, slot_id)

    @persona_group.command(name="activate", description="Set one of your personas as the active one by its slot number")
    @app_commands.describe(slot_id=f"The slot number (1-{config.MAX_PERSONAS}) of the persona to activate.")
    async def persona_activate_reg(interaction: discord.Interaction, slot_id: app_commands.Range[int, 1, config.MAX_PERSONAS]):
        """Registration wrapper for the /persona activate command."""
        await persona_activate_impl(interaction, slot_id)

    @persona_group.command(name="chat", description="Chat with your currently active custom persona")
    @app_commands.describe(message="Your message to your active persona.")
    async def persona_chat_reg(interaction: discord.Interaction, message: str):
        """Registration wrapper for the /persona chat command."""
        await persona_chat_impl(interaction, message)

    @persona_group.command(name="debate", description="Make two of your personas debate a topic")
    @app_commands.describe(slot_id_1=f"The slot number (1-{config.MAX_PERSONAS}) of the first debating persona.", slot_id_2=f"The slot number (1-{config.MAX_PERSONAS}) of the second debating persona.", topic="The topic the personas should debate.")
    async def persona_debate_reg(interaction: discord.Interaction, slot_id_1: app_commands.Range[int, 1, config.MAX_PERSONAS], slot_id_2: app_commands.Range[int, 1, config.MAX_PERSONAS], topic: str):
        """Registration wrapper for the /persona debate command."""
        await persona_debate_impl(interaction, slot_id_1, slot_id_2, topic)

    # Document Assistance Commands (Grouped)
    @doc_assist_group.command(name="analyze", description="Get an initial AI analysis of your document.")
    @app_commands.choices(doc_type=[ # Define choices for the doc_type parameter
        app_commands.Choice(name="Resume", value="resume"),
        app_commands.Choice(name="Cover Letter", value="cover_letter"),
        app_commands.Choice(name="Essay", value="essay"),
        app_commands.Choice(name="Report", value="report"),
    ])
    @app_commands.describe(doc_type="The type of document.", file="The document file to analyze.")
    async def doc_assist_analyze_reg(interaction: Interaction, doc_type: Literal['resume', 'cover_letter', 'essay', 'report'], file: Attachment):
        """Registration wrapper for the /doc_assist analyze command."""
        await doc_assist_analyze_impl(interaction, doc_type, file)

    @doc_assist_group.command(name="critique", description="Get targeted feedback on your document.")
    @app_commands.choices(doc_type=[
        app_commands.Choice(name="Resume", value="resume"),
        app_commands.Choice(name="Cover Letter", value="cover_letter"),
        app_commands.Choice(name="Essay", value="essay"),
        app_commands.Choice(name="Report", value="report"),
    ])
    @app_commands.choices(focus=[ # Define choices for the focus parameter
        app_commands.Choice(name="Clarity", value="clarity"),
        app_commands.Choice(name="Grammar & Spelling", value="grammar"),
        app_commands.Choice(name="Impact & Persuasiveness", value="impact"),
        app_commands.Choice(name="Conciseness & Flow", value="conciseness"),
    ])
    @app_commands.describe(doc_type="The type of document.", file="The document file to critique.", focus="The area to focus the critique on.")
    async def doc_assist_critique_reg(interaction: Interaction, doc_type: Literal['resume', 'cover_letter', 'essay', 'report'], file: Attachment, focus: Literal['clarity', 'grammar', 'impact', 'conciseness']):
        """Registration wrapper for the /doc_assist critique command."""
        await doc_assist_critique_impl(interaction, doc_type, file, focus)

    @doc_assist_group.command(name="rewrite_section", description="Rewrite a specific section of text for a goal.")
    @app_commands.choices(doc_type=[
        app_commands.Choice(name="Resume", value="resume"),
        app_commands.Choice(name="Cover Letter", value="cover_letter"),
        app_commands.Choice(name="Essay", value="essay"),
        app_commands.Choice(name="Report", value="report"),
    ])
    @app_commands.choices(goal=[ # Define choices for the goal parameter
        app_commands.Choice(name="Make More Professional", value="make more professional"),
        app_commands.Choice(name="Make More Concise", value="make more concise"),
        app_commands.Choice(name="Make More Engaging", value="make more engaging"),
    ])
    @app_commands.describe(doc_type="The type of document the section belongs to.", goal="The goal for rewriting.", section_text="Paste the text section to rewrite.")
    async def doc_assist_rewrite_section_reg(interaction: Interaction, doc_type: Literal['resume', 'cover_letter', 'essay', 'report'], goal: Literal['make more professional', 'make more concise', 'make more engaging'], section_text: str):
        """Registration wrapper for the /doc_assist rewrite_section command."""
        await doc_assist_rewrite_section_impl(interaction, doc_type, goal, section_text)

    # Image Generation Command
    @tree.command(name="imagine", description="Generate an image based on a text prompt using AI (Stable Diffusion).")
    @app_commands.describe(prompt="The description of the image you want to create.")
    async def imagine_command_reg(interaction: discord.Interaction, prompt: str):
        """Registration wrapper for the /imagine command."""
        # Add check for token here as well for immediate feedback if disabled.
        if not config.HUGGINGFACE_TOKEN:
            await send_error_response(interaction, "Image generation is currently disabled (missing API key).", ephemeral=True)
            return
        await imagine_command_impl(interaction, prompt)

    # Recommendation Commands (Grouped)
    @recommend_group.command(name="set_preference", description="Set your likes/dislikes for recommendation types.")
    @app_commands.choices(pref_type=[ # Define choices for preference type
        app_commands.Choice(name="Movies", value="movie"),
        app_commands.Choice(name="Books", value="book"),
        app_commands.Choice(name="Music (Artists/Genres)", value="music"),
        app_commands.Choice(name="Games", value="game"),
    ])
    @app_commands.describe(pref_type="The type of media.", likes="Comma-separated list of things you like.", dislikes="Comma-separated list of things you dislike.")
    async def recommend_set_preference_reg(interaction: Interaction, pref_type: Literal['movie', 'book', 'music', 'game'], likes: str | None = None, dislikes: str | None = None):
        """Registration wrapper for the /recommend set_preference command."""
        await recommend_set_preference_impl(interaction, pref_type, likes, dislikes)

    @recommend_group.command(name="get", description="Get personalized recommendations.")
    @app_commands.choices(rec_type=[ # Define choices for recommendation type
        app_commands.Choice(name="Movies", value="movie"),
        app_commands.Choice(name="Books", value="book"),
        app_commands.Choice(name="Music", value="music"),
        app_commands.Choice(name="Games", value="game"),
    ])
    @app_commands.describe(rec_type="What type to recommend?", genre="Optional: Filter by genre.", based_on="Optional: Recommend items similar to this (title, artist, etc.).")
    async def recommend_get_reg(interaction: Interaction, rec_type: Literal['movie', 'book', 'music', 'game'], genre: str | None = None, based_on: str | None = None):
        """Registration wrapper for the /recommend get command."""
        await recommend_get_impl(interaction, rec_type, genre, based_on)

    # Story Generator Commands (Grouped)
    @story_group.command(name="start", description="Start a new interactive story.")
    @app_commands.choices(mode=[ # Define choices for story mode
        app_commands.Choice(name="Collaborative (You & AI take turns)", value="collaborative"),
        app_commands.Choice(name="Choose Your Own (AI presents choices)", value="choose_your_own"),
    ])
    @app_commands.describe(mode="How the story progresses.", genre="Optional: Genre (e.g., fantasy, sci-fi).", setting="Optional: Brief description of the setting or starting point.")
    async def story_start_reg(interaction: Interaction, mode: Literal['collaborative', 'choose_your_own'], genre: str | None = None, setting: str | None = None):
        """Registration wrapper for the /story start command."""
        await story_start_impl(interaction, mode, genre, setting)

    @story_group.command(name="continue", description="Continue your active collaborative story.")
    @app_commands.describe(your_turn="Your contribution to the story.")
    async def story_continue_reg(interaction: Interaction, your_turn: str):
        """Registration wrapper for the /story continue command."""
        await story_continue_impl(interaction, your_turn)

    @story_group.command(name="status", description="Check the status and last turn of your active story.")
    async def story_status_reg(interaction: Interaction):
        """Registration wrapper for the /story status command."""
        await story_status_impl(interaction)

    @story_group.command(name="end", description="End your current active story session.")
    async def story_end_reg(interaction: Interaction):
        """Registration wrapper for the /story end command."""
        await story_end_impl(interaction)

    # Account/History Management Commands
    @tree.command(name="reset", description="Reset your conversation history for specific modes")
    # Dynamically create choices from VALID_CONVERSATION_TYPES in config
    @app_commands.choices(conversation_type=[
        app_commands.Choice(name="All Conversations", value="all"),
        *[app_commands.Choice(name=f"{ctype.capitalize()}", value=ctype) for ctype in config.RESETTABLE_CONVERSATION_TYPES]
    ])
    @app_commands.describe(conversation_type="Choose which conversation history to reset.")
    async def reset_command_reg(interaction: discord.Interaction, conversation_type: str):
        """Registration wrapper for the /reset command."""
        await reset_command_impl(interaction, conversation_type)

    @tree.command(name="stats", description="View your message count statistics")
    async def stats_command_reg(interaction: discord.Interaction):
        """Registration wrapper for the /stats command."""
        await stats_command_impl(interaction)

    @tree.command(name="sentiment_stats", description="Analyze sentiment of recent messages using conversational context.")
    @app_commands.choices(conversation_type=[
        app_commands.Choice(name="All Conversations", value="all"),
        # Add command name hint for clarity
        *[app_commands.Choice(name=f"{ctype.capitalize()} ({'/persona chat' if ctype == 'persona' else f'/{ctype}'})", value=ctype) for ctype in config.VALID_CONVERSATION_TYPES]
    ])
    @app_commands.describe(conversation_type="Which conversation history to analyze.", limit=f"Number of your messages to analyze (default {config.SENTIMENT_LIMIT_DEFAULT}, max {config.SENTIMENT_LIMIT_MAX}). Might be slow.")
    async def sentiment_stats_command_reg(interaction: discord.Interaction, conversation_type: str, limit: app_commands.Range[int, 1, config.SENTIMENT_LIMIT_MAX // 2] = config.SENTIMENT_LIMIT_DEFAULT // 2):
        """Registration wrapper for the /sentiment_stats command."""
        await sentiment_stats_command_impl(interaction, conversation_type, limit)

    @tree.command(name="export", description="Export your conversation history to a text file.")
    @app_commands.choices(conversation_type=[
        app_commands.Choice(name="All Conversations", value="all"),
        *[app_commands.Choice(name=f"{ctype.capitalize()}", value=ctype) for ctype in config.VALID_CONVERSATION_TYPES]
    ])
    @app_commands.describe(conversation_type="Which conversation history to export.", limit=f"Maximum number of messages to export (default {config.EXPORT_LIMIT_DEFAULT}, max {config.EXPORT_LIMIT_MAX}).")
    async def export_command_reg(interaction: discord.Interaction, conversation_type: str, limit: app_commands.Range[int, 1, config.EXPORT_LIMIT_MAX] = config.EXPORT_LIMIT_DEFAULT):
        """Registration wrapper for the /export command."""
        await export_command_impl(interaction, conversation_type, limit)

    @tree.command(name="help", description="Show available commands and bot features")
    async def help_command_reg(interaction: discord.Interaction):
        """Registration wrapper for the /help command."""
        await help_command_impl(interaction)

    # Add the command groups to the main command tree.
    tree.add_command(persona_group)
    tree.add_command(doc_assist_group)
    tree.add_command(recommend_group)
    tree.add_command(story_group)

    logger.info("Commands registered successfully.")

# END OF FILE commands.py